from flask import Flask, render_template, request, send_file, jsonify
import pandas as pd
import os
from werkzeug.utils import secure_filename
import io
import json
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch, mm
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def format_currency(value):
    """Format numbers in Indian format with commas"""
    try:
        num = float(value)
        # Format with 2 decimal places
        formatted = f"{num:.2f}"
        
        # Split into integer and decimal parts
        parts = formatted.split('.')
        integer_part = parts[0]
        decimal_part = parts[1]
        
        # Handle negative numbers
        is_negative = integer_part.startswith('-')
        if is_negative:
            integer_part = integer_part[1:]
        
        # Apply Indian numbering format
        # Last 3 digits, then every 2 digits
        if len(integer_part) <= 3:
            indian_format = integer_part
        else:
            # Reverse to process from right
            reversed_int = integer_part[::-1]
            
            # First 3 digits
            result = reversed_int[:3]
            remaining = reversed_int[3:]
            
            # Add remaining digits in groups of 2
            for i in range(0, len(remaining), 2):
                result += ',' + remaining[i:i+2]
            
            # Reverse back
            indian_format = result[::-1]
        
        # Add decimal part
        final_result = f"{indian_format}.{decimal_part}"
        
        # Add negative sign if needed
        if is_negative:
            final_result = '-' + final_result
        
        return final_result
    except:
        return value

def process_option1(input_file):
    """Filter by Hybrid Agency"""
    search_text = "TNCH00009452"
    HEADER_ROW_1 = 12
    HEADER_ROW_2 = 13
    DATA_START = 14
    HYBRID_COL_INDEX = 2

    h1 = pd.read_excel(input_file, header=None, skiprows=HEADER_ROW_1, nrows=1).iloc[0]
    h2 = pd.read_excel(input_file, header=None, skiprows=HEADER_ROW_2, nrows=1).iloc[0]
    
    h1 = h1.ffill()
    h2 = h2.fillna("")
    
    multi_columns = pd.MultiIndex.from_tuples(
        [(str(a).strip(), str(b).strip()) for a, b in zip(h1, h2)]
    )
    
    df = pd.read_excel(input_file, header=None, skiprows=DATA_START)
    df.columns = multi_columns
    df = df.ffill()
    
    filtered_df = df[
        df.iloc[:, HYBRID_COL_INDEX].astype(str).str.contains(search_text, na=False)
    ]
    
    output = io.BytesIO()
    filtered_df.to_excel(output, index=True)
    output.seek(0)
    
    return output

def extract_uc_data(input_file):
    """Extract UC data based on Hybrid Agency value - supports multiple rows"""
    try:
        search_value = "TNCH00009452"
        HEADER_ROW_1 = 12
        HEADER_ROW_2 = 13
        DATA_START = 14
        
        # Read header rows
        h1 = pd.read_excel(input_file, header=None, skiprows=HEADER_ROW_1, nrows=1).iloc[0]
        h2 = pd.read_excel(input_file, header=None, skiprows=HEADER_ROW_2, nrows=1).iloc[0]
        
        h1 = h1.ffill()
        h2 = h2.fillna("")
        
        # Create multi-index columns
        multi_columns = pd.MultiIndex.from_tuples(
            [(str(a).strip(), str(b).strip()) for a, b in zip(h1, h2)]
        )
        
        # Read data
        df = pd.read_excel(input_file, header=None, skiprows=DATA_START)
        df.columns = multi_columns
        df = df.ffill()
        
        # Find the Hybrid Agency column by checking column names
        hybrid_agency_col = None
        for col in df.columns:
            col_str = str(col)
            if 'Hybrid Agency' in col_str:
                hybrid_agency_col = col
                break
        
        if hybrid_agency_col is None:
            print("Warning: Hybrid Agency column not found")
            print(f"Available columns: {df.columns.tolist()}")
            return None
        
        # Find rows matching the search value
        # Convert to Series and then apply the filter
        hybrid_agency_series = df[hybrid_agency_col]
        if isinstance(hybrid_agency_series, pd.DataFrame):
            # If multi-column result, take first column
            hybrid_agency_series = hybrid_agency_series.iloc[:, 0]
        
        filtered_df = df[hybrid_agency_series.astype(str).str.contains(search_value, na=False)]
        
        if filtered_df.empty:
            print(f"Warning: No rows found with Hybrid Agency value: {search_value}")
            return None
        
        # Find the Assignment Sanction Number column
        assignment_sanction_col = None
        for col in df.columns:
            if 'Assignment Sanction' in str(col):
                assignment_sanction_col = col
                break
        
        # Find the Total Drawing limit issued by Parent Agency column
        total_drawing_col = None
        for col in df.columns:
            if 'Total Drawing limit issued by Parent Agency' in str(col):
                total_drawing_col = col
                break
        
        # Find columns for the last three columns (indices 6, 7, 8)
        total_available_col = None
        expenditure_col = None
        closing_balance_col = None
        
        for col in df.columns:
            col_str = str(col)
            if 'Total Available funds' in col_str:
                total_available_col = col
            elif 'Expenditure incurred' in col_str:
                expenditure_col = col
            elif 'Closing Balances' in col_str:
                closing_balance_col = col
        
        print(f"DEBUG: Assignment Sanction Col: {assignment_sanction_col}")
        print(f"DEBUG: Total Drawing Col: {total_drawing_col}")
        print(f"DEBUG: Total Available Col: {total_available_col}")
        print(f"DEBUG: Expenditure Col: {expenditure_col}")
        print(f"DEBUG: Closing Balance Col: {closing_balance_col}")
        print(f"DEBUG: Found {len(filtered_df)} matching rows")
        
        # Extract data from ALL matching rows
        uc_data_list = []
        seen_combinations = set()  # Track seen sanction+amount combinations to avoid duplicates
        
        for idx, (_, row) in enumerate(filtered_df.iterrows()):
            sanction_number = ""
            amount = 0
            
            # Extract Sanction Number (take the part ending with (G) or (C))
            if assignment_sanction_col is not None:
                try:
                    sanction_value = row[assignment_sanction_col]
                    # Handle if it returns a DataFrame/Series
                    if isinstance(sanction_value, pd.Series):
                        sanction_value = sanction_value.iloc[0] if len(sanction_value) > 0 else ""
                    
                    sanction_str = str(sanction_value).strip()
                    print(f"DEBUG Row {idx}: Sanction String: {sanction_str}")
                    
                    # Extract the part ending with (G) or (C) - handle both uppercase and lowercase
                    # Look for the rightmost occurrence of (G), (g), (C), or (c)
                    sanction_number = ""
                    
                    # Find all positions of suffixes
                    pos_G = sanction_str.rfind('(G)')
                    pos_g = sanction_str.rfind('(g)')
                    pos_C = sanction_str.rfind('(C)')
                    pos_c = sanction_str.rfind('(c)')
                    
                    # Find the rightmost position
                    positions = [(pos, 3) for pos in [pos_G, pos_g, pos_C, pos_c] if pos != -1]
                    
                    if positions:
                        # Get the rightmost position
                        max_pos = max(positions, key=lambda x: x[0])[0]
                        sanction_number = sanction_str[:max_pos + 3]
                    else:
                        sanction_number = sanction_str
                    
                    print(f"DEBUG Row {idx}: Extracted Sanction Number: {sanction_number}")
                except Exception as e:
                    print(f"Error extracting sanction number for row {idx}: {e}")
                    sanction_number = ""
            
            # Extract Total Drawing limit
            if total_drawing_col is not None:
                try:
                    amount_value = row[total_drawing_col]
                    # Handle if it returns a DataFrame/Series
                    if isinstance(amount_value, pd.Series):
                        amount_value = amount_value.iloc[0] if len(amount_value) > 0 else 0
                    
                    amount = float(amount_value)
                    print(f"DEBUG Row {idx}: Amount: {amount}")
                except Exception as e:
                    print(f"Error extracting amount for row {idx}: {e}")
                    amount = 0
            
            # Extract Total Available funds
            total_available = 0
            if total_available_col is not None:
                try:
                    val = row[total_available_col]
                    if isinstance(val, pd.Series):
                        val = val.iloc[0] if len(val) > 0 else 0
                    total_available = float(val)
                except Exception as e:
                    total_available = 0
            
            # Extract Expenditure incurred
            expenditure = 0
            if expenditure_col is not None:
                try:
                    val = row[expenditure_col]
                    if isinstance(val, pd.Series):
                        val = val.iloc[0] if len(val) > 0 else 0
                    expenditure = float(val)
                except Exception as e:
                    expenditure = 0
            
            # Extract Closing Balances
            closing_balance = 0
            if closing_balance_col is not None:
                try:
                    val = row[closing_balance_col]
                    if isinstance(val, pd.Series):
                        val = val.iloc[0] if len(val) > 0 else 0
                    closing_balance = float(val)
                except Exception as e:
                    closing_balance = 0
            
            # Only add if we have at least a sanction number and it's not a duplicate
            if sanction_number:
                # Round amount to 2 decimal places to avoid floating point precision issues
                # This ensures "104730990.0" and "104730990.00" are treated the same
                amount_rounded = round(amount, 2)
                combo = (sanction_number.upper(), amount_rounded)  # Normalize sanction number to uppercase for comparison
                
                if combo not in seen_combinations:
                    uc_data_list.append({
                        'sanction_number': sanction_number,
                        'amount': amount,
                        'total_available': total_available,
                        'expenditure': expenditure,
                        'closing_balance': closing_balance
                    })
                    seen_combinations.add(combo)
                    print(f"DEBUG: Added unique entry - Sanction: {sanction_number}, Amount: {amount_rounded}")
                else:
                    print(f"DEBUG: Skipped duplicate entry - Sanction: {sanction_number}, Amount: {amount_rounded}")
        
        if not uc_data_list:
            print("No valid UC data extracted")
            return None
        
        # Separate data by grant type (G = Recurring, C = Non-Recurring)
        recurring_data = []
        non_recurring_data = []
        
        for entry in uc_data_list:
            sanction_num = entry['sanction_number'].upper().strip()
            # Check suffix while handling any whitespace
            if sanction_num.endswith('(G)') or '(G)' in sanction_num[-5:]:
                recurring_data.append(entry)
                print(f"DEBUG: Added to Recurring - {entry['sanction_number']}")
            elif sanction_num.endswith('(C)') or '(C)' in sanction_num[-5:]:
                non_recurring_data.append(entry)
                print(f"DEBUG: Added to Non-Recurring - {entry['sanction_number']}")
            else:
                # If no clear suffix, log it
                print(f"DEBUG: WARNING - No clear suffix detected for {entry['sanction_number']}, sanction_num={sanction_num}")
        
        print(f"DEBUG: Extracted {len(uc_data_list)} unique UC data entries (after removing duplicates)")
        print(f"DEBUG: Recurring (G) entries: {len(recurring_data)}")
        print(f"DEBUG: Non-Recurring (C) entries: {len(non_recurring_data)}")
        print(f"DEBUG: Recurring data: {[(d['sanction_number'], d['amount']) for d in recurring_data]}")
        print(f"DEBUG: Non-Recurring data: {[(d['sanction_number'], d['amount']) for d in non_recurring_data]}")
        
        # Return organized by type
        return {
            'recurring': recurring_data,
            'non_recurring': non_recurring_data
        }
    
    except Exception as e:
        print(f"Error extracting UC data: {e}")
        import traceback
        traceback.print_exc()
        return None

def extract_dates(input_file):
    """Extract From Date and To Date from Excel file metadata"""
    try:
        metadata = pd.read_excel(input_file, header=None, nrows=20)
        
        from_date = None
        to_date = None
        
        print(f"\n=== DEBUG: Extracting Dates ===")
        
        # Search for "From Date" and "To Date"
        for row_idx, row in metadata.iterrows():
            row_str = ' '.join([str(x) for x in row if x is not None and x != ''])
            
            # Check for From Date
            if "from date" in row_str.lower():
                print(f"Found From Date reference at row {row_idx}")
                for col_idx, value in enumerate(row):
                    if value is None or value == '':
                        continue
                    
                    str_value = str(value).strip()
                    
                    if "from date" in str_value.lower():
                        # Look in next columns for the date
                        for offset in range(1, min(5, metadata.shape[1] - col_idx)):
                            next_val = row.iloc[col_idx + offset] if col_idx + offset < len(row) else None
                            if next_val is not None and next_val != '':
                                next_str = str(next_val).strip()
                                # Check if it looks like a date
                                if any(c.isdigit() for c in next_str) and len(next_str) >= 8:
                                    from_date = next_str
                                    print(f"Found From Date: {from_date}")
                                    break
            
            # Check for To Date
            if "to date" in row_str.lower():
                print(f"Found To Date reference at row {row_idx}")
                for col_idx, value in enumerate(row):
                    if value is None or value == '':
                        continue
                    
                    str_value = str(value).strip()
                    
                    if "to date" in str_value.lower():
                        # Look in next columns for the date
                        for offset in range(1, min(5, metadata.shape[1] - col_idx)):
                            next_val = row.iloc[col_idx + offset] if col_idx + offset < len(row) else None
                            if next_val is not None and next_val != '':
                                next_str = str(next_val).strip()
                                # Check if it looks like a date
                                if any(c.isdigit() for c in next_str) and len(next_str) >= 8:
                                    to_date = next_str
                                    print(f"Found To Date: {to_date}")
                                    break
        
        return from_date, to_date
    except Exception as e:
        print(f"Error extracting dates: {e}")
        import traceback
        traceback.print_exc()
        return None, None

def extract_financial_year(input_file):
    """Extract Financial Year from Excel file metadata"""
    try:
        # Read the first 20 rows without header to get metadata
        metadata = pd.read_excel(input_file, header=None, nrows=20)
        
        print(f"\n=== DEBUG: Extracting Financial Year from {input_file} ===")
        print(f"Metadata shape: {metadata.shape}")
        
        # Search for "Financial Year" in the first 20 rows
        for row_idx, row in metadata.iterrows():
            row_str = ' '.join([str(x) for x in row if x is not None and x != ''])
            
            # Check if this row contains financial year info
            if "financial year" in row_str.lower():
                print(f"Found Financial Year reference at row {row_idx}")
                print(f"Row content: {row_str}")
                
                for col_idx, value in enumerate(row):
                    if value is None or value == '':
                        continue
                    
                    str_value = str(value).strip()
                    
                    # Check if current cell contains both label and date
                    if "financial year" in str_value.lower() and ("(" in str_value or "-" in str_value):
                        # Extract just the date part
                        if "(" in str_value:
                            date_part = str_value[str_value.find("("):]
                            print(f"Found combined label+date at row {row_idx}, col {col_idx}: {date_part}")
                            return date_part
                    
                    # Check next columns for date range
                    if "financial year" in str_value.lower():
                        # Look in subsequent columns for the date
                        for offset in range(1, min(5, metadata.shape[1] - col_idx)):
                            next_val = row.iloc[col_idx + offset] if col_idx + offset < len(row) else None
                            if next_val is not None and next_val != '':
                                next_str = str(next_val).strip()
                                # Check if it looks like a date range
                                if "(" in next_str or ("-" in next_str and any(c.isdigit() for c in next_str)):
                                    print(f"Found date at row {row_idx}, col {col_idx + offset}: {next_str}")
                                    return next_str
        
        print("Financial Year not found in metadata")
        return None
    except Exception as e:
        print(f"Error extracting financial year: {e}")
        import traceback
        traceback.print_exc()
        return None

def process_option2(input_file, original_file=None):
    """Child Agency Summary"""
    df = pd.read_excel(input_file, header=[0, 1])
    
    df.columns = [
        " ".join([str(a), str(b)]).replace("nan", "").strip()
        for a, b in df.columns
    ]
    
    thub_cols = {"self_limit": None, "self_success": None, "self_pending": None}
    
    for c in df.columns:
        cl = c.lower()
        if "amount kept for self expenditure" in cl:
            thub_cols["self_limit"] = c
        elif "total expenditure incurred" in cl and "success" in cl and "by self" in cl:
            thub_cols["self_success"] = c
        elif "total expenditure incurred" in cl and "pending" in cl and "by self" in cl:
            thub_cols["self_pending"] = c
    
    required_cols = {
        "child_agency": None, "sanction": None, "total": None,
        "success": None, "pending": None
    }
    
    for c in df.columns:
        cl = c.lower()
        if "child" in cl and "agency" in cl:
            required_cols["child_agency"] = c
        elif "assignment sanction" in cl:
            required_cols["sanction"] = c
        elif "child expenditure limit assigned" in cl:
            required_cols["total"] = c
        elif "success" in cl and "child agencies" in cl:
            required_cols["success"] = c
        elif "pending" in cl and "child agencies" in cl:
            required_cols["pending"] = c
    
    child_agency_col = required_cols["child_agency"]
    sanction_col = required_cols["sanction"]
    total_col = required_cols["total"]
    success_col = required_cols["success"]
    pending_col = required_cols["pending"]
    
    df[sanction_col] = df[sanction_col].astype(str).str.replace(r"\s+\((C|G)\)", r"(\1)", regex=True)
    
    def grant_type(val):
        val = str(val).upper()
        if val.endswith("(G)"):
            return "Recurring"
        if val.endswith("(C)"):
            return "Non-Recurring"
        return "Unknown"
    
    df["Grant Type"] = df[sanction_col].apply(grant_type)
    
    for c in [total_col, success_col, pending_col]:
        df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0)
    
    for key in thub_cols:
        if thub_cols[key]:
            df[thub_cols[key]] = pd.to_numeric(df[thub_cols[key]], errors="coerce").fillna(0)
    
    df["T-Hub Expenditure Spent"] = df[thub_cols["self_success"]] + df[thub_cols["self_pending"]]
    df["T-Hub Balance"] = df[thub_cols["self_limit"]] - df["T-Hub Expenditure Spent"]
    df["Spent"] = df[success_col] + df[pending_col]
    
    summary = df.groupby([child_agency_col, sanction_col, "Grant Type"], dropna=False).agg(
        Total=(total_col, "sum"),
        Spent=("Spent", "sum")
    ).reset_index()
    
    summary["Balance"] = summary["Total"] - summary["Spent"]
    
    summary = summary.rename(columns={
        child_agency_col: "Child Agency Name",
        sanction_col: "Assignment Sanction Number",
        "Total": "Expenditure Limit",
        "Spent": "Expenditure Spent"
    })
    
    summary = summary[
        summary["Child Agency Name"].notna() &
        (summary["Child Agency Name"].astype(str).str.strip() != "")
    ]
    
    summary["Child Agency Name"] = summary["Child Agency Name"].replace({
        "[BRPA00011340] - IIT Patna NQM": "IITP",
        "[CDACT] - C-DAC,Trivandrum": "CDAC-T",
        "[CDOT] - CENTRE FOR DEVELOPMENT OF TELEMATICS": "CDOT",
        "[CTDU00009549] - IIT Bhilai NQM": "IITBhilai",
        "[DLND00006296] - IITD (IRD) Delhi NQM": "IITD",
        "[IITTP] - Indian Institute of Technology Tirupati": "IITTP",
        "[JKJA00011940] - IIT Jammu NQM": "IITJammu",
        "[KABN00017992] - IISc Bangalore NQM": "IIScB",
        "[KABN00018025] - C-DAC BLR NQM": "CDAC-B",
        "[KABN00018031] - Raman Research Institute QC NQM": "RRI",
        "[MPBP00008046] - IISER Bhopal NQM": "IISERB",
        "[MPIN00008812] - IIT Indore NQM": "IITI",
        "[PBSA00009654] - IISER Mohali NQM": "IISERMohali",
        "[TLHY00007631] - IIT hyderabad NQM": "IITH",
        "[TNCH00009491] - INDIAN INSTITUTE OF TECHNOLOGY, MADRAS_3": "IITM",
        "[TNCH00009619] - Society for Electronic Transactions and Security (SETS) NQM": "SETS",
        "[UKHA00008653] - IIT ROORKEE NQM": "IITR",
        "[UPKS00018833] - IIT Kanpur NQM": "IITK",
        "[WBME00009648] - IIT Kharagpur NQM": "IITKgp",
        "[UPAH00015132] - HRI - NQM": "HRI",
        "[KABN00018141] - IISc. Bangalore NQM (T3)": "IISc-TG3"
    })
    
    def normalize_name(val):
        return str(val).upper().replace(" ", "").replace("-", "").replace(",", "")
    
    TG_MAP = {
        "IITP": "TG4", "CDACT": "TG1", "CDOT": "TG2", "IITBHILAI": "TG2",
        "IITD": "TG3", "IITTP": "TG4", "IITJAMMU": "TG1", "IISCB": "TG1",
        "CDACB": "TG1", "RRI": "TG4", "IISERB": "TG3", "IITI": "TG1",
        "IISERMOHALI": "TG4", "IITH": "TG1", "IITM": "TG1", "SETS": "TG1",
        "IITR": "TG3", "IITK": "TG2", "IITKGP": "TG2", "HRI": "TG4", "IISCTG3": "TG3"
    }
    
    summary["TG"] = summary["Child Agency Name"].apply(
        lambda x: next((tg for key, tg in TG_MAP.items() if key in normalize_name(x)), "")
    )
    
    summary = summary[["TG", "Child Agency Name", "Assignment Sanction Number", 
                       "Grant Type", "Expenditure Limit", "Expenditure Spent", "Balance"]]
    summary = summary.sort_values(by=["TG", "Child Agency Name", "Assignment Sanction Number"])
    
    rn_summary = summary.groupby("Grant Type", dropna=False).agg(
        Total_Expenditure_Limit=("Expenditure Limit", "sum"),
        Total_Expenditure_Spent=("Expenditure Spent", "sum"),
        Total_Balance=("Balance", "sum")
    ).reset_index()
    
    inst_rn = summary.pivot_table(
        index="Child Agency Name",
        columns="Grant Type",
        values=["Expenditure Limit", "Expenditure Spent", "Balance"],
        aggfunc="sum",
        fill_value=0
    )
    inst_rn.columns = [f"{gt} - {metric}" for metric, gt in inst_rn.columns]
    inst_rn = inst_rn.reset_index()
    inst_rn["TG"] = inst_rn["Child Agency Name"].apply(
        lambda x: next((tg for key, tg in TG_MAP.items() if key in normalize_name(x)), "")
    )
    
    # Sort by TG and Child Agency Name first
    inst_rn = inst_rn.sort_values(by=["TG", "Child Agency Name"]).reset_index(drop=True)
    
    # Reorder columns: TG, Child Agency Name, Recurring (all 3), Non-Recurring (all 3)
    col_order = ["TG", "Child Agency Name"]
    
    # Add Recurring columns in order
    for metric in ["Expenditure Limit", "Expenditure Spent", "Balance"]:
        col_name = f"Recurring - {metric}"
        if col_name in inst_rn.columns:
            col_order.append(col_name)
    
    # Add Non-Recurring columns in order
    for metric in ["Expenditure Limit", "Expenditure Spent", "Balance"]:
        col_name = f"Non-Recurring - {metric}"
        if col_name in inst_rn.columns:
            col_order.append(col_name)
    
    # Select only columns in the desired order
    inst_rn = inst_rn[[col for col in col_order if col in inst_rn.columns]]
    
    thub_summary = df.groupby([sanction_col, "Grant Type"], dropna=False).agg(
        Expenditure_Limit=(thub_cols["self_limit"], "max"),
        Expenditure_Spent=("T-Hub Expenditure Spent", "max")
    ).reset_index()
    thub_summary["Balance"] = thub_summary["Expenditure_Limit"] - thub_summary["Expenditure_Spent"]
    thub_summary = thub_summary.rename(columns={sanction_col: "Assignment Sanction Number"})
    thub_summary = thub_summary[
        thub_summary["Assignment Sanction Number"].notna() &
        (thub_summary["Assignment Sanction Number"].astype(str).str.strip() != "") &
        (thub_summary["Grant Type"] != "Unknown")
    ]
    thub_summary.insert(0, "Hub", "Samgnya")
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        workbook = writer.book
        
        # Number and font formats
        num_format = workbook.add_format({'num_format': '#,##0.00', 'font_name': 'Segoe UI', 'font_size': 11, 'align': 'left', 'valign': 'vcenter'})
        body_format = workbook.add_format({'font_name': 'Open Sans', 'font_size': 11, 'align': 'left', 'valign': 'vcenter'})
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': 'white',
            'font_color': 'black',
            'align': 'center',
            'valign': 'vcenter',
            'border': 1,
            'font_name': 'Open Sans',
            'font_size': 11
        })

        # Write sheets in correct order: Institute, T-Hub, Overall
        # Sheet 1: Institute-Exp Rec vs Non-Rec
        inst_rn.to_excel(writer, sheet_name="Institute-Exp Rec vs Non-Rec", index=False)
        worksheet = writer.sheets["Institute-Exp Rec vs Non-Rec"]
        worksheet.set_column('A:Z', 14)  # Set all columns to 14 width for screen fit
        # Format header row
        for col_num, col_name in enumerate(inst_rn.columns):
            worksheet.write(0, col_num, col_name, header_format)
        # Apply formatting to all data cells
        for row_num, row_data in enumerate(inst_rn.values):
            for col_num, value in enumerate(row_data):
                if col_num < len(inst_rn.columns):
                    col_name = inst_rn.columns[col_num]
                    if any(x in col_name for x in ["Expenditure", "Balance"]):
                        worksheet.write(row_num + 1, col_num, value, num_format)
                    else:
                        worksheet.write(row_num + 1, col_num, value, body_format)

        # Sheet 2: TG-Wise Exp Recurring
        tg_summary_for_excel = inst_rn.groupby('TG', as_index=False).agg({
            col: 'sum' for col in inst_rn.columns if col not in ['TG', 'Child Agency Name']
        })
        cols = ['TG'] + [col for col in tg_summary_for_excel.columns if col != 'TG']
        tg_summary_for_excel = tg_summary_for_excel[cols]
        grand_total_tg = {col: tg_summary_for_excel[col].sum() for col in tg_summary_for_excel.columns if col != 'TG'}
        grand_total_tg['TG'] = 'Grand Total'
        tg_summary_for_excel = pd.concat([tg_summary_for_excel, pd.DataFrame([grand_total_tg])], ignore_index=True)
        
        tg_summary_for_excel.to_excel(writer, sheet_name="TG-Wise Exp Rec vs Non-Rec", index=False)
        worksheet = writer.sheets["TG-Wise Exp Rec vs Non-Rec"]
        worksheet.set_column('A:Z', 14)
        # Format header row
        for col_num, col_name in enumerate(tg_summary_for_excel.columns):
            worksheet.write(0, col_num, col_name, header_format)
        # Apply formatting to all data cells
        for row_num, row_data in enumerate(tg_summary_for_excel.values):
            for col_num, value in enumerate(row_data):
                if col_num < len(tg_summary_for_excel.columns):
                    col_name = tg_summary_for_excel.columns[col_num]
                    if any(x in col_name for x in ["Expenditure", "Balance"]):
                        worksheet.write(row_num + 1, col_num, value, num_format)
                    else:
                        worksheet.write(row_num + 1, col_num, value, body_format)

        # Sheet 3: T-Hub-Exp Rec vs Non-Rec
        thub_summary.to_excel(writer, sheet_name="T-Hub-Exp Rec vs Non-Rec", index=False)
        worksheet = writer.sheets["T-Hub-Exp Rec vs Non-Rec"]
        worksheet.set_column('A:Z', 15)  # Set all columns to 15 width for screen fit
        # Format header row
        for col_num, col_name in enumerate(thub_summary.columns):
            worksheet.write(0, col_num, col_name, header_format)
        # Apply formatting to all data cells
        for row_num, row_data in enumerate(thub_summary.values):
            for col_num, value in enumerate(row_data):
                if col_num < len(thub_summary.columns):
                    col_name = thub_summary.columns[col_num]
                    if col_name in ["Expenditure_Limit", "Expenditure_Spent", "Balance"]:
                        worksheet.write(row_num + 1, col_num, value, num_format)
                    else:
                        worksheet.write(row_num + 1, col_num, value, body_format)

        # Sheet 4: Overall Rec vs Non-Rec
        rn_summary.to_excel(writer, sheet_name="Overall Rec vs Non-Rec", index=False)
        worksheet = writer.sheets["Overall Rec vs Non-Rec"]
        worksheet.set_column('A:Z', 18)  # Wider columns for 4-column sheet
        # Format header row with blue background
        for col_num, col_name in enumerate(rn_summary.columns):
            worksheet.write(0, col_num, col_name, header_format)
        # Apply formatting to all data cells
        for row_num, row_data in enumerate(rn_summary.values):
            for col_num, value in enumerate(row_data):
                if col_num < len(rn_summary.columns):
                    col_name = rn_summary.columns[col_num]
                    if "Total" in col_name or "Balance" in col_name:
                        worksheet.write(row_num + 1, col_num, value, num_format)
                    else:
                        worksheet.write(row_num + 1, col_num, value, body_format)
    
    output.seek(0)
    
    # Create TG-wise summary table from inst_rn data
    # Group by TG and sum all columns
    tg_summary_table = inst_rn.groupby('TG', as_index=False).agg({
        col: 'sum' for col in inst_rn.columns if col not in ['TG', 'Child Agency Name']
    })
    
    # Reorder columns: TG first, then all other columns
    cols = ['TG'] + [col for col in tg_summary_table.columns if col != 'TG']
    tg_summary_table = tg_summary_table[cols]
    
    # Calculate Grand Total row
    grand_total = {col: tg_summary_table[col].sum() for col in tg_summary_table.columns if col != 'TG'}
    grand_total['TG'] = 'Grand Total'
    
    # Add Grand Total row
    tg_summary_table = pd.concat([tg_summary_table, pd.DataFrame([grand_total])], ignore_index=True)
    
    # Convert to list of dictionaries for JSON serialization
    tg_summary_table_data = tg_summary_table.to_dict('records')
    
    # Create TG-wise summary with recurring data grouped by TG
    tg_wise_data = {}
    
    # Group summary data by TG
    for tg in summary['TG'].unique():
        if tg and tg != "":
            tg_data = summary[summary['TG'] == tg]
            
            # Filter only Recurring data for TG
            tg_recurring = tg_data[tg_data['Grant Type'] == 'Recurring']
            
            if len(tg_recurring) > 0:
                # Calculate totals for this TG
                total_funds_released = tg_recurring['Expenditure Limit'].sum()
                total_expenditure = tg_recurring['Expenditure Spent'].sum()
                balance = tg_recurring['Balance'].sum()
                
                tg_wise_data[tg] = {
                    'TG': tg,
                    'Total_Funds_Released': total_funds_released,
                    'Total_Expenditure': total_expenditure,
                    'Balance': balance,
                    'institutions': tg_data[['Child Agency Name', 'Assignment Sanction Number']].drop_duplicates().to_dict('records')
                }
    
    # Extract Financial Year and Dates from original file if provided
    financial_year = None
    from_date = None
    to_date = None
    
    if original_file and os.path.exists(original_file):
        financial_year = extract_financial_year(original_file)
        from_date, to_date = extract_dates(original_file)
    
    # Calculate T-Hub Totals (sum of only Recurring Grant Type)
    thub_recurring = thub_summary[thub_summary['Grant Type'] == 'Recurring']
    thub_totals = {
        'total_funds_released': thub_recurring['Expenditure_Limit'].sum(),
        'total_expenditure': thub_recurring['Expenditure_Spent'].sum(),
        'balance': thub_recurring['Balance'].sum()
    }
    
    # Return both file and preview data (removed 'institute_wise')
    return output, {
        'recurring_summary': rn_summary.to_dict('records'),
        'institute_split': inst_rn.to_dict('records'),
        'tg_summary_table': tg_summary_table_data,
        'thub_summary': thub_summary.to_dict('records'),
        'thub_totals': thub_totals,
        'tg_wise_summary': tg_wise_data,
        'financial_year': financial_year,
        'from_date': from_date,
        'to_date': to_date
    }

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    filepath = None
    intermediate_path = None
    
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not file.filename.endswith(('.xlsx', '.xls')):
            return jsonify({'error': 'Please upload an Excel file'}), 400
        
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        print("Processing Step 1: Filtering data...")
        intermediate_output = process_option1(filepath)
        
        intermediate_path = os.path.join(app.config['UPLOAD_FOLDER'], 'intermediate_filtered.xlsx')
        with open(intermediate_path, 'wb') as f:
            f.write(intermediate_output.getvalue())
        
        print("Processing Step 2: Creating child agency summary...")
        final_output, preview_data = process_option2(intermediate_path, filepath)
        
        # Save final output temporarily for download
        final_path = os.path.join(app.config['UPLOAD_FOLDER'], 'final_output.xlsx')
        with open(final_path, 'wb') as f:
            f.write(final_output.getvalue())
        
        # Extract UC data from the original input file
        print("Processing Step 3: Extracting UC data...")
        uc_data = extract_uc_data(filepath)
        if uc_data:
            preview_data['uc_data'] = uc_data
            print(f"UC data extracted: {uc_data}")
        else:
            preview_data['uc_data'] = None
        
        if filepath and os.path.exists(filepath):
            os.remove(filepath)
        if intermediate_path and os.path.exists(intermediate_path):
            os.remove(intermediate_path)
        
        return jsonify({
            'success': True,
            'preview': preview_data,
            'download_url': '/download'
        })
    
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print("ERROR:", error_trace)
        
        if filepath and os.path.exists(filepath):
            os.remove(filepath)
        if intermediate_path and os.path.exists(intermediate_path):
            os.remove(intermediate_path)
        
        return jsonify({'error': str(e), 'traceback': error_trace}), 500

@app.route('/download-tg-pdf', methods=['POST'])
def download_tg_pdf():
    """Generate and download TG template as PDF"""
    try:
        data = request.get_json()
        tg = data.get('tg')
        total_funds_released = data.get('total_funds_released', 0)
        total_expenditure = data.get('total_expenditure', 0)
        balance = data.get('balance', 0)
        to_date = data.get('to_date', '')
        
        # Create PDF
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=A4, rightMargin=15*mm, leftMargin=15*mm, topMargin=20*mm, bottomMargin=20*mm)
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#0066cc'),
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        label_style = ParagraphStyle(
            'Label',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#495057'),
            spaceAfter=10
        )
        
        # Build document
        story = []
        
        # Title
        story.append(Paragraph(f"{tg} - Institutions Expenditure Summary", title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Date
        if to_date:
            story.append(Paragraph(f"<b>Balance as on:</b> {to_date}", label_style))
            story.append(Spacer(1, 0.2*inch))
        
        # Create header style
        header_cell_style = ParagraphStyle(
            'HeaderCell',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.HexColor('#000000'),
            alignment=TA_CENTER,
            spaceAfter=0,
            leading=10
        )
        
        # Table data with Paragraph headers for proper wrapping
        table_data = [
            [
                Paragraph('Sanctioned Head', header_cell_style),
                Paragraph('Total Funds<br/>Released', header_cell_style),
                Paragraph('Total<br/>Expenditure', header_cell_style),
                Paragraph('Balance as on<br/>(DD/MM/YYYY)', header_cell_style),
                Paragraph('Remarks<br/>(if any)', header_cell_style)
            ],
            ['(I)', '(II)', '(III)', '(IV)', ''],
            ['Recurring', '', '', '', '(*)'],
            ['Total', 
             f"{total_funds_released:,.2f}", 
             f"{total_expenditure:,.2f}", 
             f"{balance:,.2f}", 
             '']
        ]
        
        # Create table with optimized column widths
        table = Table(table_data, colWidths=[1.3*inch, 1.2*inch, 1.2*inch, 1.3*inch, 0.9*inch], rowHeights=[0.6*inch, 0.25*inch, 0.25*inch, 0.25*inch])
        
        table.setStyle(TableStyle([
            # Header row with word wrapping
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFFFFF')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#000000')),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('LEFTPADDING', (0, 0), (-1, 0), 6),
            ('RIGHTPADDING', (0, 0), (-1, 0), 6),
            ('TOPPADDING', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('WORDWRAP', (0, 0), (-1, 0), True),
            
            # Subheader row
            ('FONTNAME', (0, 1), (-1, 1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 1), (-1, 1), 9),
            ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor('#e3f2fd')),
            ('ALIGN', (0, 1), (-1, 1), 'CENTER'),
            ('VALIGN', (0, 1), (-1, 1), 'MIDDLE'),
            ('TOPPADDING', (0, 1), (-1, 1), 6),
            ('BOTTOMPADDING', (0, 1), (-1, 1), 6),
            
            # Data rows
            ('FONTNAME', (0, 2), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 2), (-1, -1), 9),
            ('ALIGN', (0, 2), (-1, -1), 'CENTER'),
            ('ALIGN', (0, 2), (0, -1), 'LEFT'),
            ('VALIGN', (0, 2), (-1, -1), 'MIDDLE'),
            ('ROWBACKGROUNDS', (0, 2), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#dee2e6')),
            ('LEFTPADDING', (0, 2), (-1, -1), 6),
            ('RIGHTPADDING', (0, 2), (-1, -1), 6),
            ('TOPPADDING', (0, 2), (-1, -1), 7),
            ('BOTTOMPADDING', (0, 2), (-1, -1), 7),
        ]))
        
        story.append(table)
        
        # Build PDF
        doc.build(story)
        pdf_buffer.seek(0)
        
        return send_file(
            pdf_buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{tg}_expenditure_summary.pdf'
        )
    
    except Exception as e:
        print(f"Error generating PDF: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/download')
def download():
    final_path = os.path.join(app.config['UPLOAD_FOLDER'], 'final_output.xlsx')
    
    if not os.path.exists(final_path):
        return jsonify({'error': 'File not found'}), 404
    
    return send_file(
        final_path,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name='child_agency_summary.xlsx'
    )

@app.route('/download-tg-tables', methods=['POST'])
def download_tg_tables():
    """Download TG or UC tables in selected format (Excel, PDF, or Word)"""
    try:
        data = request.get_json()
        format_type = data.get('format', 'excel')
        table_data = data.get('data', [])
        sheet_type = data.get('sheetType', 'all')  # 'recurring', 'nonrecurring', or 'all'
        is_uc = data.get('isUC', False)  # Flag to indicate if this is UC data
        
        if not table_data:
            return jsonify({'error': 'No data provided'}), 400
        
        # Determine filename based on sheet type and whether it's UC or TG
        if is_uc:
            if sheet_type == 'recurring':
                base_filename = 'UC_Exp_Recurring'
            elif sheet_type == 'nonrecurring':
                base_filename = 'UC_Exp_Non-Recurring'
            else:
                base_filename = 'All_UC_Detailed_Tables'
        else:
            if sheet_type == 'recurring':
                base_filename = 'TG-Wise_Exp_Recurring'
            elif sheet_type == 'nonrecurring':
                base_filename = 'TG-Wise_Exp_Non-Recurring'
            else:
                base_filename = 'All_TG_Detailed_Tables'
        
        if format_type == 'excel':
            if is_uc:
                output = generate_uc_excel(table_data)
            else:
                output = generate_tg_excel(table_data)
            filename = f'{base_filename}.xlsx'
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif format_type == 'pdf':
            if is_uc:
                output = generate_uc_pdf(table_data)
            else:
                output = generate_tg_pdf(table_data)
            filename = f'{base_filename}.pdf'
            mimetype = 'application/pdf'
        elif format_type == 'word':
            if is_uc:
                output = generate_uc_word(table_data)
            else:
                output = generate_tg_word(table_data)
            filename = f'{base_filename}.docx'
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            return jsonify({'error': 'Invalid format'}), 400
        
        return send_file(
            output,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        print(f"Error downloading tables: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# UC Excel, PDF, and Word generation functions (simplified - reuse TG functions for now)
def generate_uc_excel(uc_data):
    """Generate UC tables in Excel format"""
    # For now, we can create a simple excel with the UC data
    # This converts the UC data to a format similar to TG data
    wb = Workbook()
    ws = wb.active
    ws.title = 'UC Data'
    
    # Add headers
    if isinstance(uc_data, dict) and uc_data.get('ucTables'):
        uc_tables = uc_data.get('ucTables', [])
        
        # If it's a list of UC tables, create sheets for each
        if isinstance(uc_tables, list) and len(uc_tables) > 0:
            wb.remove(ws)
            
            for idx, uc_table in enumerate(uc_tables):
                ws = wb.create_sheet(f'UC Data {idx + 1}' if idx > 0 else 'UC Data')
                
                if isinstance(uc_table, dict):
                    # Add table data
                    row = 1
                    for key, value in uc_table.items():
                        ws.cell(row=row, column=1).value = str(key)
                        ws.cell(row=row, column=2).value = str(value)
                        row += 1
    
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def generate_uc_pdf(uc_data):
    """Generate UC tables in PDF format"""
    # Create a simple PDF with UC data
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    c.setTitle("UC Data Export")
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, "UC Utilization Certificate Data")
    
    c.setFont("Helvetica", 12)
    y_position = height - 100
    
    if isinstance(uc_data, dict) and uc_data.get('ucTables'):
        uc_tables = uc_data.get('ucTables', [])
        
        if isinstance(uc_tables, list):
            for idx, uc_table in enumerate(uc_tables):
                if isinstance(uc_table, dict):
                    c.drawString(50, y_position, f"UC Table {idx + 1}:")
                    y_position -= 20
                    
                    for key, value in uc_table.items():
                        c.drawString(70, y_position, f"{key}: {value}")
                        y_position -= 15
                        
                        if y_position < 50:
                            c.showPage()
                            c.setFont("Helvetica", 12)
                            y_position = height - 50
                    
                    y_position -= 20
    
    c.save()
    buffer.seek(0)
    return buffer

def generate_uc_word(uc_data):
    """Generate UC tables in Word format"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('UC Utilization Certificate Data', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if isinstance(uc_data, dict) and uc_data.get('ucTables'):
        uc_tables = uc_data.get('ucTables', [])
        
        if isinstance(uc_tables, list):
            for idx, uc_table in enumerate(uc_tables):
                doc.add_heading(f'UC Table {idx + 1}', level=2)
                
                if isinstance(uc_table, dict):
                    table = doc.add_table(rows=len(uc_table) + 1, cols=2)
                    table.style = 'Light Grid Accent 1'
                    
                    # Add headers
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Field'
                    hdr_cells[1].text = 'Value'
                    
                    # Add data
                    for row_idx, (key, value) in enumerate(uc_table.items(), 1):
                        cells = table.rows[row_idx].cells
                        cells[0].text = str(key)
                        cells[1].text = str(value)
                
                doc.add_paragraph()  # Add spacing between tables
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

@app.route('/download-thub-tables', methods=['POST'])
def download_thub_tables():
    """Download T-Hub tables in selected format (Excel, PDF, or Word)"""
    try:
        data = request.get_json()
        format_type = data.get('format', 'excel')
        thub_data = data.get('data', [])
        
        if not thub_data:
            return jsonify({'error': 'No T-Hub data provided'}), 400
        
        if format_type == 'excel':
            output = generate_thub_excel(thub_data)
            filename = 'All_T-Hub_Tables.xlsx'
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif format_type == 'pdf':
            output = generate_thub_pdf(thub_data)
            filename = 'All_T-Hub_Tables.pdf'
            mimetype = 'application/pdf'
        elif format_type == 'word':
            output = generate_thub_word(thub_data)
            filename = 'All_T-Hub_Tables.docx'
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            return jsonify({'error': 'Invalid format'}), 400
        
        return send_file(
            output,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        print(f"Error downloading T-Hub tables: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

def generate_thub_excel(thub_data):
    """Generate T-Hub table in Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = 'T-Hub'
    
    # Define border style
    thin_border = Border(
        left=Side(style='thin', color='CCCCCC'),
        right=Side(style='thin', color='CCCCCC'),
        top=Side(style='thin', color='CCCCCC'),
        bottom=Side(style='thin', color='CCCCCC')
    )
    
    # Add title
    ws['A1'] = 'T-Hub - Expenditure Summary'
    ws['A1'].font = Font(size=14, bold=True, color='FFFFFF')
    ws['A1'].fill = PatternFill(start_color='0066CC', end_color='0066CC', fill_type='solid')
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.merge_cells('A1:F1')
    ws.row_dimensions[1].height = 25
    
    # Add headers
    headers = ['Hub', 'Assignment Sanction Number', 'Grant Type', 'Total Funds Released\n(II)', 'Total Expenditure\n(III)', 'Balance\n(II - III)']
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col_num)
        cell.value = header
        cell.font = Font(bold=True, color='000000', size=12)
        cell.fill = PatternFill(start_color='FFFFFF', end_color='FFFFFF', fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = Border(
            left=Side(style='medium', color='000000'),
            right=Side(style='medium', color='000000'),
            top=Side(style='medium', color='000000'),
            bottom=Side(style='medium', color='000000')
        )
    ws.row_dimensions[3].height = 40
    
    # Add data
    for row_num, row_data in enumerate(thub_data, 4):
        ws.cell(row=row_num, column=1).value = row_data.get('hub', '')
        ws.cell(row=row_num, column=2).value = row_data.get('sanction_number', '')
        ws.cell(row=row_num, column=3).value = row_data.get('grant_type', '')
        ws.cell(row=row_num, column=4).value = row_data.get('total_funds_released', 0)
        ws.cell(row=row_num, column=5).value = row_data.get('total_expenditure', 0)
        ws.cell(row=row_num, column=6).value = row_data.get('balance', 0)
        
        # Format numeric columns and add borders
        for col in [1, 2, 3, 4, 5, 6]:
            cell = ws.cell(row=row_num, column=col)
            cell.border = thin_border
            cell.alignment = Alignment(vertical='center')
            if col in [4, 5, 6]:
                cell.number_format = '#,##0.00'
                cell.alignment = Alignment(horizontal='right', vertical='center')
            else:
                cell.alignment = Alignment(horizontal='left', vertical='center')
    
    # Set column widths
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 25
    ws.column_dimensions['E'].width = 25
    ws.column_dimensions['F'].width = 20
    
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def generate_thub_pdf(thub_data):
    """Generate T-Hub table in PDF format"""
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=A4, rightMargin=12*mm, leftMargin=12*mm, topMargin=15*mm, bottomMargin=15*mm)
    story = []
    
    # Styles
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'Title',
        fontSize=16,
        textColor=colors.HexColor('#0066CC'),
        spaceAfter=20,
        alignment=1,  # CENTER
        fontName='Helvetica-Bold'
    )
    title = Paragraph('T-Hub - Expenditure Summary', title_style)
    story.append(title)
    
    # Create header style for proper text wrapping
    header_cell_style = ParagraphStyle(
        'HeaderCell',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.HexColor('#000000'),
        alignment=TA_CENTER,
        spaceAfter=0,
        leading=10
    )
    
    # Table data with Paragraph headers
    table_data = [
        [
            Paragraph('Hub', header_cell_style),
            Paragraph('Assignment<br/>Sanction Number', header_cell_style),
            Paragraph('Grant<br/>Type', header_cell_style),
            Paragraph('Total Funds<br/>Released (II)', header_cell_style),
            Paragraph('Total<br/>Expenditure (III)', header_cell_style),
            Paragraph('Balance<br/>(II - III)', header_cell_style)
        ]
    ]
    for row in thub_data:
        table_data.append([
            row.get('hub', ''),
            row.get('sanction_number', ''),
            row.get('grant_type', ''),
            f"{float(row.get('total_funds_released', 0)):,.2f}",
            f"{float(row.get('total_expenditure', 0)):,.2f}",
            f"{float(row.get('balance', 0)):,.2f}"
        ])
    
    # Create table with optimized column widths for A4
    table = Table(table_data, colWidths=[0.9*inch, 1.3*inch, 1.0*inch, 1.2*inch, 1.2*inch, 1.1*inch], rowHeights=[0.5*inch, None])
    table.setStyle(TableStyle([
        # Header row styling with word wrapping
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFFFFF')),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#000000')),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('LEFTPADDING', (0, 0), (-1, 0), 5),
        ('RIGHTPADDING', (0, 0), (-1, 0), 5),
        ('WORDWRAP', (0, 0), (-1, 0), True),
        
        # Data rows styling
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ('LEFTPADDING', (0, 1), (-1, -1), 5),
        ('RIGHTPADDING', (0, 1), (-1, -1), 5),
        
        # Grid and row styling
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CCCCCC')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8F9FA')]),
    ]))
    
    story.append(table)
    doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer

def generate_thub_word(thub_data):
    """Generate T-Hub table in Word format"""
    doc = Document()
    
    # Add title
    title = doc.add_paragraph('T-Hub - Expenditure Summary')
    title.style = 'Heading 1'
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if title.runs:
        title_format = title.runs[0]
        title_format.font.size = Pt(14)
        title_format.font.color.rgb = RGBColor(0, 102, 204)
        title_format.font.bold = True
    
    # Add table
    headers = ['Hub', 'Assignment Sanction Number', 'Grant Type', 'Total Funds Released\n(II)', 'Total Expenditure\n(III)', 'Balance\n(II - III)']
    
    rows = len(thub_data) + 1
    table = doc.add_table(rows=rows, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        # Set header cell background color
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading_elm = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))
        header_cells[idx]._element.get_or_add_tcPr().append(shading_elm)
        
        # Clear cell and set text with line breaks
        header_cells[idx].text = ''
        paragraph = header_cells[idx].paragraphs[0]
        
        # Split header by newlines and add each line as a separate run
        lines = header.split('\n')
        for line_idx, line in enumerate(lines):
            if line_idx > 0:
                paragraph.add_run('\n')
            run = paragraph.add_run(line)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(12)
        
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data rows
    for row_idx, row_data in enumerate(thub_data, 1):
        cells = table.rows[row_idx].cells
        cells[0].text = row_data.get('hub', '')
        cells[1].text = row_data.get('sanction_number', '')
        cells[2].text = row_data.get('grant_type', '')
        cells[3].text = f"{float(row_data.get('total_funds_released', 0)):,.2f}"
        cells[4].text = f"{float(row_data.get('total_expenditure', 0)):,.2f}"
        cells[5].text = f"{float(row_data.get('balance', 0)):,.2f}"
        
        # Center align numeric columns
        for col_idx in [3, 4, 5]:
            for paragraph in cells[col_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def generate_tg_excel(tg_data):
    """Generate TG tables in Excel format with separate sheets for each table"""
    wb = Workbook()
    wb.remove(wb.active)
    
    # Define border style
    thin_border = Border(
        left=Side(style='thin', color='CCCCCC'),
        right=Side(style='thin', color='CCCCCC'),
        top=Side(style='thin', color='CCCCCC'),
        bottom=Side(style='thin', color='CCCCCC')
    )
    
    # Check if tg_data is the new format with separate table types
    if isinstance(tg_data, dict):
        # New format with comparisonTable, thubSummaryTable, tgDetailedTables
        
        # Create Comparison Table sheet
        if tg_data.get('comparisonTable') and len(tg_data['comparisonTable']) > 0:
            comp_table = tg_data['comparisonTable'][0]
            ws = wb.create_sheet('Total Expenditure (T-Hub & TGs)')
            add_comparison_table_to_sheet(ws, comp_table, thin_border)
        
        # Create T-Hub Summary Table sheet
        if tg_data.get('thubSummaryTable') and len(tg_data['thubSummaryTable']) > 0:
            thub_summary = tg_data['thubSummaryTable'][0]
            ws = wb.create_sheet('T-Hub-Wise Expenditure Summary')
            add_thub_summary_table_to_sheet(ws, thub_summary, thin_border)
        
        # Create TG Detailed Tables sheets
        if tg_data.get('tgDetailedTables'):
            for tg_item in tg_data['tgDetailedTables']:
                ws = wb.create_sheet(tg_item.get('tgName', 'TG'))
                add_tg_detail_table_to_sheet(ws, tg_item, thin_border)
    else:
        # Old format - backward compatibility
        for tg_item in tg_data:
            ws = wb.create_sheet(tg_item.get('tgName', 'TG'))
            add_tg_detail_table_to_sheet(ws, tg_item, thin_border)
    
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def add_comparison_table_to_sheet(ws, comp_table, thin_border):
    """Add comparison table to a worksheet"""
    # Add title
    ws['A1'] = 'Total Expenditure (T-Hub & TGs)'
    ws['A1'].font = Font(size=14, bold=True, color='FFFFFF')
    ws['A1'].fill = PatternFill(start_color='0066CC', end_color='0066CC', fill_type='solid')
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.merge_cells('A1:D1')
    ws.row_dimensions[1].height = 25
    
    # Add headers
    headers = ['T-Hub & TG\n(I)', 'Total Funds Released\n(II)', 'Total Expenditure\n(III)', 'Balance\n(IV = II - III)', 'Remarks\n(if any)']
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col_num)
        cell.value = header
        cell.font = Font(bold=True, color='000000', size=12)
        cell.fill = PatternFill(start_color='FFFFFF', end_color='FFFFFF', fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = Border(
            left=Side(style='medium', color='000000'),
            right=Side(style='medium', color='000000'),
            top=Side(style='medium', color='000000'),
            bottom=Side(style='medium', color='000000')
        )
    ws.row_dimensions[3].height = 40
    
    # Add data rows
    for row_num, row_data in enumerate(comp_table.get('rows', []), 4):
        ws.cell(row=row_num, column=1).value = row_data.get('name', '')
        
        try:
            ws.cell(row=row_num, column=2).value = float(row_data.get('fundsReleased', 0)) if row_data.get('fundsReleased', 0) else 0
            ws.cell(row=row_num, column=3).value = float(row_data.get('expenditure', 0)) if row_data.get('expenditure', 0) else 0
            ws.cell(row=row_num, column=4).value = float(row_data.get('balance', 0)) if row_data.get('balance', 0) else 0
        except (ValueError, TypeError):
            ws.cell(row=row_num, column=2).value = 0
            ws.cell(row=row_num, column=3).value = 0
            ws.cell(row=row_num, column=4).value = 0
        
        ws.cell(row=row_num, column=5).value = row_data.get('remarks', '')
        
        for col in [1, 2, 3, 4, 5]:
            cell = ws.cell(row=row_num, column=col)
            cell.border = thin_border
            if col in [2, 3, 4]:
                cell.number_format = '#,##0.00'
                cell.alignment = Alignment(horizontal='right', vertical='center')
            else:
                cell.alignment = Alignment(horizontal='left', vertical='center')
    
    # Set column widths
    ws.column_dimensions['A'].width = 22
    ws.column_dimensions['B'].width = 28
    ws.column_dimensions['C'].width = 28
    ws.column_dimensions['D'].width = 28
    ws.column_dimensions['E'].width = 22

def add_thub_summary_table_to_sheet(ws, thub_summary, thin_border):
    """Add T-Hub summary table to a worksheet"""
    # Add title
    ws['A1'] = 'T-Hub-Wise Expenditure Summary'
    ws['A1'].font = Font(size=14, bold=True, color='FFFFFF')
    ws['A1'].fill = PatternFill(start_color='0066CC', end_color='0066CC', fill_type='solid')
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.merge_cells('A1:E1')
    ws.row_dimensions[1].height = 25
    
    # Add headers
    headers = ['Sanctioned Head\n(I)', 'Total Funds Released\n(II)', 'Total Expenditure\n(III)', 'Balance\n(IV = II - III)', 'Remarks\n(if any)']
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col_num)
        cell.value = header
        cell.font = Font(bold=True, color='000000', size=12)
        cell.fill = PatternFill(start_color='FFFFFF', end_color='FFFFFF', fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = Border(
            left=Side(style='medium', color='000000'),
            right=Side(style='medium', color='000000'),
            top=Side(style='medium', color='000000'),
            bottom=Side(style='medium', color='000000')
        )
    ws.row_dimensions[3].height = 40
    
    # Add data rows
    for row_num, row_data in enumerate(thub_summary.get('rows', []), 4):
        ws.cell(row=row_num, column=1).value = row_data.get('sanctioned_head', '')
        
        try:
            ws.cell(row=row_num, column=2).value = float(row_data.get('total_funds_released', 0)) if row_data.get('total_funds_released', 0) else 0
            ws.cell(row=row_num, column=3).value = float(row_data.get('total_expenditure', 0)) if row_data.get('total_expenditure', 0) else 0
            ws.cell(row=row_num, column=4).value = float(row_data.get('balance', 0)) if row_data.get('balance', 0) else 0
        except (ValueError, TypeError):
            ws.cell(row=row_num, column=2).value = 0
            ws.cell(row=row_num, column=3).value = 0
            ws.cell(row=row_num, column=4).value = 0
        
        ws.cell(row=row_num, column=5).value = row_data.get('remarks', '')
        
        for col in [1, 2, 3, 4, 5]:
            cell = ws.cell(row=row_num, column=col)
            cell.border = thin_border
            if col in [2, 3, 4]:
                cell.number_format = '#,##0.00'
                cell.alignment = Alignment(horizontal='right', vertical='center')
            else:
                cell.alignment = Alignment(horizontal='left', vertical='center')
    
    # Set column widths
    ws.column_dimensions['A'].width = 22
    ws.column_dimensions['B'].width = 28
    ws.column_dimensions['C'].width = 28
    ws.column_dimensions['D'].width = 28
    ws.column_dimensions['E'].width = 22

def add_tg_detail_table_to_sheet(ws, tg_item, thin_border):
    """Add TG detail table to a worksheet"""
    # Add title
    ws['A1'] = f"{tg_item.get('tgName')} - {tg_item.get('institutionName')}"
    ws['A1'].font = Font(size=14, bold=True, color='FFFFFF')
    ws['A1'].fill = PatternFill(start_color='0066CC', end_color='0066CC', fill_type='solid')
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.merge_cells('A1:E1')
    ws.row_dimensions[1].height = 25
    
    # Add headers
    headers = ['Sanctioned Head\n(I)', 'Total Funds Released\n(II)', 'Total Expenditure\n(III)', 'Balance as on\n(IV = II - III)', 'Remarks\n(if any)']

    to_date = tg_item.get('toDate', '30-12-2025')
    headers[3] = f'Balance as on ({to_date})\n(IV = II - III)'
    
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col_num)
        cell.value = header
        cell.font = Font(bold=True, color='000000', size=12)
        cell.fill = PatternFill(start_color='FFFFFF', end_color='FFFFFF', fill_type='solid')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = Border(
            left=Side(style='medium', color='000000'),
            right=Side(style='medium', color='000000'),
            top=Side(style='medium', color='000000'),
            bottom=Side(style='medium', color='000000')
        )
    ws.row_dimensions[3].height = 40
    
    # Add data
    for row_num, row_data in enumerate(tg_item.get('rows', []), 4):
        ws.cell(row=row_num, column=1).value = row_data.get('sanctioned_head', '')
        
        # Convert numeric values to float to ensure proper formatting
        try:
            ws.cell(row=row_num, column=2).value = float(row_data.get('total_funds_released', 0)) if row_data.get('total_funds_released', 0) else 0
            ws.cell(row=row_num, column=3).value = float(row_data.get('total_expenditure', 0)) if row_data.get('total_expenditure', 0) else 0
            ws.cell(row=row_num, column=4).value = float(row_data.get('balance', 0)) if row_data.get('balance', 0) else 0
        except (ValueError, TypeError):
            ws.cell(row=row_num, column=2).value = 0
            ws.cell(row=row_num, column=3).value = 0
            ws.cell(row=row_num, column=4).value = 0
        
        ws.cell(row=row_num, column=5).value = row_data.get('remarks', '')
        
        # Format numeric columns and add borders
        for col in [1, 2, 3, 4, 5]:
            cell = ws.cell(row=row_num, column=col)
            cell.border = thin_border
            cell.alignment = Alignment(vertical='center')
            if col in [2, 3, 4]:
                cell.number_format = '#,##0.00'
                cell.alignment = Alignment(horizontal='right', vertical='center')
            else:
                cell.alignment = Alignment(horizontal='left', vertical='center')
    
    # Set column widths
    ws.column_dimensions['A'].width = 22
    ws.column_dimensions['B'].width = 28
    ws.column_dimensions['C'].width = 28
    ws.column_dimensions['D'].width = 40
    ws.column_dimensions['E'].width = 22

def generate_tg_pdf(tg_data):
    """Generate TG tables in PDF format"""
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
    story = []
    
    # Check if new format
    if isinstance(tg_data, dict):
        # Add Comparison Table
        if tg_data.get('comparisonTable') and len(tg_data['comparisonTable']) > 0:
            comp_table = tg_data['comparisonTable'][0]
            add_comparison_table_to_pdf(story, comp_table)
            story.append(Spacer(1, 0.08*inch))
        
        # Add T-Hub Summary Table
        if tg_data.get('thubSummaryTable') and len(tg_data['thubSummaryTable']) > 0:
            thub_summary = tg_data['thubSummaryTable'][0]
            add_thub_summary_table_to_pdf(story, thub_summary)
            story.append(Spacer(1, 0.08*inch))
        
        # Add TG Detailed Tables
        if tg_data.get('tgDetailedTables'):
            for idx, tg_item in enumerate(tg_data['tgDetailedTables']):
                add_tg_detail_table_to_pdf(story, tg_item)
                if idx < len(tg_data['tgDetailedTables']) - 1:
                    story.append(Spacer(1, 0.08*inch))
    else:
        # Old format
        for idx, tg_item in enumerate(tg_data):
            if idx > 0:
                story.append(Spacer(1, 0.08*inch))
            add_tg_detail_table_to_pdf(story, tg_item)
    
    doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer

def add_comparison_table_to_pdf(story, comp_table):
    """Add comparison table to PDF story"""
    title_style = ParagraphStyle(
        'Title',
        fontSize=14,
        textColor=colors.HexColor('#0066CC'),
        spaceAfter=6,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )
    title = Paragraph('Total Expenditure (T-Hub & TGs)', title_style)
    story.append(title)
    
    header_style = ParagraphStyle('HeaderStyle', fontSize=9, fontName='Helvetica-Bold', alignment=TA_CENTER)
    headers = [
        Paragraph('T-Hub & TG<br/>(I)', header_style),
        Paragraph('Total Funds Released<br/>(II)', header_style),
        Paragraph('Total Expenditure<br/>(III)', header_style),
        Paragraph('Balance<br/>(IV = II - III)', header_style),
        Paragraph('Remarks<br/>(if any)', header_style)
    ]
    
    table_data = [headers]
    for row in comp_table.get('rows', []):
        try:
            funds = float(row.get('fundsReleased', 0)) if row.get('fundsReleased', 0) else 0
            expenditure = float(row.get('expenditure', 0)) if row.get('expenditure', 0) else 0
            balance = float(row.get('balance', 0)) if row.get('balance', 0) else 0
        except (ValueError, TypeError):
            funds = 0
            expenditure = 0
            balance = 0
        
        table_data.append([
            row.get('name', ''),
            format_currency(funds),
            format_currency(expenditure),
            format_currency(balance),
            row.get('remarks', '')
        ])
    
    table = Table(table_data, colWidths=[1.2*inch, 1.4*inch, 1.4*inch, 1.4*inch, 1.0*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFFFFF')),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 8),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CCCCCC')),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.08*inch))

def add_thub_summary_table_to_pdf(story, thub_summary):
    """Add T-Hub summary table to PDF story"""
    title_style = ParagraphStyle(
        'Title',
        fontSize=14,
        textColor=colors.HexColor('#0066CC'),
        spaceAfter=6,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )
    title = Paragraph('T-Hub-Wise Expenditure Summary', title_style)
    story.append(title)
    
    header_style = ParagraphStyle('HeaderStyle', fontSize=9, fontName='Helvetica-Bold', alignment=TA_CENTER)
    headers = [
        Paragraph('Sanctioned Head<br/>(I)', header_style),
        Paragraph('Total Funds Released<br/>(II)', header_style),
        Paragraph('Total Expenditure<br/>(III)', header_style),
        Paragraph('Balance<br/>(IV = II - III)', header_style),
        Paragraph('Remarks<br/>(if any)', header_style)
    ]
    
    table_data = [headers]
    for row in thub_summary.get('rows', []):
        try:
            funds = float(row.get('total_funds_released', 0)) if row.get('total_funds_released', 0) else 0
            expenditure = float(row.get('total_expenditure', 0)) if row.get('total_expenditure', 0) else 0
            balance = float(row.get('balance', 0)) if row.get('balance', 0) else 0
        except (ValueError, TypeError):
            funds = 0
            expenditure = 0
            balance = 0
        
        table_data.append([
            row.get('sanctioned_head', ''),
            format_currency(funds),
            format_currency(expenditure),
            format_currency(balance),
            row.get('remarks', '')
        ])
    
    table = Table(table_data, colWidths=[1.2*inch, 1.4*inch, 1.4*inch, 1.4*inch, 1.0*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFFFFF')),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 8),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CCCCCC')),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.08*inch))

def add_tg_detail_table_to_pdf(story, tg_item):
    """Add TG detail table to PDF story"""
    title_style = ParagraphStyle(
        'TGTitle',
        fontSize=14,
        textColor=colors.HexColor('#0066CC'),
        spaceAfter=6,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )
    title = Paragraph(f"{tg_item.get('tgName')} - {tg_item.get('institutionName')}", title_style)
    story.append(title)
    
    to_date = tg_item.get('toDate', '30-12-2025')
    header_style = ParagraphStyle('HeaderStyle', fontSize=9, fontName='Helvetica-Bold', alignment=TA_CENTER)
    headers = [
        Paragraph('Sanctioned Head<br/>(I)', header_style),
        Paragraph('Total Funds Released<br/>(II)', header_style),
        Paragraph('Total Expenditure<br/>(III)', header_style),
        Paragraph(f'Balance as on ({to_date})<br/>(IV = II - III)', header_style),
        Paragraph('Remarks<br/>(if any)', header_style)
    ]
    
    table_data = [headers]
    for row in tg_item.get('rows', []):
        try:
            funds = float(row.get('total_funds_released', 0)) if row.get('total_funds_released', 0) else 0
            expenditure = float(row.get('total_expenditure', 0)) if row.get('total_expenditure', 0) else 0
            balance = float(row.get('balance', 0)) if row.get('balance', 0) else 0
        except (ValueError, TypeError):
            funds = 0
            expenditure = 0
            balance = 0
        
        table_data.append([
            row.get('sanctioned_head', ''),
            format_currency(funds),
            format_currency(expenditure),
            format_currency(balance),
            row.get('remarks', '')
        ])
    
    table = Table(table_data, colWidths=[1.2*inch, 1.4*inch, 1.4*inch, 1.8*inch, 1.0*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFFFFF')),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 8),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CCCCCC')),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.08*inch))

def generate_tg_word(tg_data):
    """Generate TG tables in Word format with separate sections for each table"""
    doc = Document()
    
    # Check if new format
    if isinstance(tg_data, dict):
        # Add Comparison Table
        if tg_data.get('comparisonTable') and len(tg_data['comparisonTable']) > 0:
            comp_table = tg_data['comparisonTable'][0]
            add_comparison_table_to_word(doc, comp_table)
        
        # Add T-Hub Summary Table
        if tg_data.get('thubSummaryTable') and len(tg_data['thubSummaryTable']) > 0:
            thub_summary = tg_data['thubSummaryTable'][0]
            add_thub_summary_table_to_word(doc, thub_summary)
        
        # Add TG Detailed Tables
        if tg_data.get('tgDetailedTables'):
            for tg_item in tg_data['tgDetailedTables']:
                add_tg_detail_table_to_word(doc, tg_item)
    else:
        # Old format - backward compatibility
        for tg_item in tg_data:
            add_tg_detail_table_to_word(doc, tg_item)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def set_table_black_borders(table):
    """Set all table borders to normal black color"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        tbl.insert(0, tblPr)
    
    # Remove existing borders if any
    existing_borders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    
    # Create table borders element with normal black color (000000) and reduced thickness
    tblBorders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000" w:themeColor="none"/>'
        r'<w:left w:val="single" w:sz="12" w:space="0" w:color="000000" w:themeColor="none"/>'
        r'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000" w:themeColor="none"/>'
        r'<w:right w:val="single" w:sz="12" w:space="0" w:color="000000" w:themeColor="none"/>'
        r'<w:insideH w:val="single" w:sz="12" w:space="0" w:color="000000" w:themeColor="none"/>'
        r'<w:insideV w:val="single" w:sz="12" w:space="0" w:color="000000" w:themeColor="none"/>'
        r'</w:tblBorders>'
    )
    
    tblPr.append(tblBorders)

def add_comparison_table_to_word(doc, comp_table):
    """Add comparison table to Word document"""
    # Add title
    title = doc.add_paragraph('Total Expenditure (T-Hub & TGs)')
    title.style = 'Heading 2'
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        title.runs[0].font.bold = True
    
    headers = ['T-Hub & TG\n(I)', 'Total Funds Released\n(II)', 'Total Expenditure\n(III)', 
              'Balance\n(IV = II - III)', 'Remarks\n(if any)']
    
    rows = len(comp_table.get('rows', [])) + 1
    table = doc.add_table(rows=rows, cols=5)
    table.style = 'Table Grid'
    set_table_black_borders(table)
    
    # Add headers
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading_elm = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))
        header_cells[idx]._element.get_or_add_tcPr().append(shading_elm)
        
        header_cells[idx].text = ''
        paragraph = header_cells[idx].paragraphs[0]
        lines = header.split('\n')
        for line_idx, line in enumerate(lines):
            if line_idx > 0:
                paragraph.add_run('\n')
            run = paragraph.add_run(line)
            run.font.bold = True
            run.font.size = Pt(11)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data rows
    for row_idx, row_data in enumerate(comp_table.get('rows', []), 1):
        cells = table.rows[row_idx].cells
        cells[0].text = row_data.get('name', '')
        
        try:
            funds = float(row_data.get('fundsReleased', 0)) if row_data.get('fundsReleased', 0) else 0
            expenditure = float(row_data.get('expenditure', 0)) if row_data.get('expenditure', 0) else 0
            balance = float(row_data.get('balance', 0)) if row_data.get('balance', 0) else 0
        except (ValueError, TypeError):
            funds = 0
            expenditure = 0
            balance = 0
        
        cells[1].text = format_currency(funds)
        cells[2].text = format_currency(expenditure)
        cells[3].text = format_currency(balance)
        cells[4].text = row_data.get('remarks', '')
        
        for col_idx in [1, 2, 3]:
            for paragraph in cells[col_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_thub_summary_table_to_word(doc, thub_summary):
    """Add T-Hub summary table to Word document"""
    # Add title
    title = doc.add_paragraph('T-Hub-Wise Expenditure Summary')
    title.style = 'Heading 2'
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        title.runs[0].font.bold = True
    
    headers = ['Sanctioned Head\n(I)', 'Total Funds Released\n(II)', 'Total Expenditure\n(III)', 
              'Balance\n(IV = II - III)', 'Remarks\n(if any)']
    
    rows = len(thub_summary.get('rows', [])) + 1
    table = doc.add_table(rows=rows, cols=5)
    table.style = 'Table Grid'
    set_table_black_borders(table)
    
    # Add headers
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading_elm = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))
        header_cells[idx]._element.get_or_add_tcPr().append(shading_elm)
        
        header_cells[idx].text = ''
        paragraph = header_cells[idx].paragraphs[0]
        lines = header.split('\n')
        for line_idx, line in enumerate(lines):
            if line_idx > 0:
                paragraph.add_run('\n')
            run = paragraph.add_run(line)
            run.font.bold = True
            run.font.size = Pt(11)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data rows
    for row_idx, row_data in enumerate(thub_summary.get('rows', []), 1):
        cells = table.rows[row_idx].cells
        cells[0].text = row_data.get('sanctioned_head', '')
        
        try:
            funds = float(row_data.get('total_funds_released', 0)) if row_data.get('total_funds_released', 0) else 0
            expenditure = float(row_data.get('total_expenditure', 0)) if row_data.get('total_expenditure', 0) else 0
            balance = float(row_data.get('balance', 0)) if row_data.get('balance', 0) else 0
        except (ValueError, TypeError):
            funds = 0
            expenditure = 0
            balance = 0
        
        cells[1].text = format_currency(funds)
        cells[2].text = format_currency(expenditure)
        cells[3].text = format_currency(balance)
        cells[4].text = row_data.get('remarks', '')
        
        for col_idx in [1, 2, 3]:
            for paragraph in cells[col_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_tg_detail_table_to_word(doc, tg_item):
    """Add TG detail table to Word document"""
    # Add title
    title = doc.add_paragraph(f"{tg_item.get('tgName')} - {tg_item.get('institutionName')}")
    title.style = 'Heading 2'
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        title.runs[0].font.bold = True
    
    # Add table
    to_date = tg_item.get('toDate', '30-12-2025')
    headers = ['Sanctioned Head\n(I)', 'Total Funds Released\n(II)', 'Total Expenditure\n(III)', 
              f'Balance as on ({to_date})\n(IV = II - III)', 'Remarks\n(if any)']
    
    rows = len(tg_item.get('rows', [])) + 1
    table = doc.add_table(rows=rows, cols=5)
    table.style = 'Table Grid'
    set_table_black_borders(table)
    
    # Add headers
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        # Remove header cell background color (white)
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading_elm = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))
        header_cells[idx]._element.get_or_add_tcPr().append(shading_elm)
        
        # Clear cell and set text with line breaks
        header_cells[idx].text = ''
        paragraph = header_cells[idx].paragraphs[0]
        
        # Split header by newlines and add each line as a separate run
        lines = header.split('\n')
        for line_idx, line in enumerate(lines):
            if line_idx > 0:
                paragraph.add_run('\n')
            run = paragraph.add_run(line)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
            run.font.size = Pt(11)
        
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data rows
    for row_idx, row_data in enumerate(tg_item.get('rows', []), 1):
        cells = table.rows[row_idx].cells
        cells[0].text = row_data.get('sanctioned_head', '')
        
        try:
            funds_released = float(row_data.get('total_funds_released', 0)) if row_data.get('total_funds_released', 0) else 0
            expenditure = float(row_data.get('total_expenditure', 0)) if row_data.get('total_expenditure', 0) else 0
            balance = float(row_data.get('balance', 0)) if row_data.get('balance', 0) else 0
        except (ValueError, TypeError):
            funds_released = 0
            expenditure = 0
            balance = 0
        
        cells[1].text = format_currency(funds_released)
        cells[2].text = format_currency(expenditure)
        cells[3].text = format_currency(balance)
        cells[4].text = row_data.get('remarks', '')
        
        # Center align numeric columns
        for col_idx in [1, 2, 3]:
            for paragraph in cells[col_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

@app.route('/download-thub-totals', methods=['POST'])
def download_thub_totals():
    """Download T-Hub Totals in selected format (Excel, PDF, or Word)"""
    try:
        data = request.get_json()
        format_type = data.get('format', 'excel')
        thub_data = data.get('data', {})
        to_date = data.get('toDate', '')
        
        if not thub_data:
            return jsonify({'error': 'No T-Hub data provided'}), 400
        
        if format_type == 'excel':
            output = generate_thub_totals_excel(thub_data, to_date)
            filename = 'T-Hub_Expenditure_Totals.xlsx'
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif format_type == 'pdf':
            output = generate_thub_totals_pdf(thub_data, to_date)
            filename = 'T-Hub_Expenditure_Totals.pdf'
            mimetype = 'application/pdf'
        elif format_type == 'word':
            output = generate_thub_totals_word(thub_data, to_date)
            filename = 'T-Hub_Expenditure_Totals.docx'
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            return jsonify({'error': 'Invalid format'}), 400
        
        return send_file(
            output,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        print(f"Error downloading T-Hub totals: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

def generate_thub_totals_excel(thub_data, to_date):
    """Generate T-Hub Totals in Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = "T-Hub Totals"
    
    # Add title
    ws['A1'] = "T-Hub-Wise Expenditure Summary"
    ws['A1'].font = Font(bold=True, size=14)
    
    # Add headers
    headers = ['Sanctioned Head (I)', 'Total Funds Released (II)', 'Total Expenditure (III)', 
               f'Balance as on ({to_date}) (IV = II - III)', 'Remarks (if any)']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col)
        cell.value = header
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="343A40", end_color="343A40", fill_type="solid")
    
    # Add data rows
    recurring_row = [
        'Recurring',
        thub_data.get('total_funds_released', 0),
        thub_data.get('total_expenditure', 0),
        thub_data.get('balance', 0),
        ''
    ]
    
    total_row = [
        'Total',
        thub_data.get('total_funds_released', 0),
        thub_data.get('total_expenditure', 0),
        thub_data.get('balance', 0),
        ''
    ]
    
    for row_idx, row_data in enumerate([recurring_row, total_row], 4):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            if col_idx == 1:
                cell.value = value
            else:
                if isinstance(value, (int, float)) and value != '':
                    cell.value = float(value)
                    cell.number_format = '#,##0.00'
                else:
                    cell.value = value
    
    # Format Total row
    for col in range(1, 6):
        ws.cell(row=5, column=col).font = Font(bold=True)
        ws.cell(row=5, column=col).fill = PatternFill(start_color="E3F2FD", end_color="E3F2FD", fill_type="solid")
    
    # Set column widths
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 25
    ws.column_dimensions['D'].width = 30
    ws.column_dimensions['E'].width = 20
    
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def generate_thub_totals_pdf(thub_data, to_date):
    """Generate T-Hub Totals in PDF format"""
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch, 
                           leftMargin=0.5*inch, rightMargin=0.5*inch)
    elements = []
    
    # Add title
    title_style = ParagraphStyle(
        'Title',
        fontSize=16,
        textColor=colors.HexColor('#0066CC'),
        spaceAfter=20,
        alignment=1,  # CENTER
        fontName='Helvetica-Bold'
    )
    title = Paragraph("T-Hub-Wise Expenditure Summary", title_style)
    elements.append(title)
    
    # Create header style with word wrapping
    header_style = ParagraphStyle(
        'HeaderStyle',
        fontSize=9,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#000000'),
        alignment=1,  # CENTER
        wordWrap='CJK',
        leading=12
    )
    
    # Create headers with proper wrapping
    headers = [
        Paragraph('Sanctioned Head<br/>(I)', header_style),
        Paragraph('Total Funds Released<br/>(II)', header_style),
        Paragraph('Total Expenditure<br/>(III)', header_style),
        Paragraph(f'Balance as on ({to_date})<br/>(IV = II - III)', header_style),
        Paragraph('Remarks<br/>(if any)', header_style)
    ]
    
    # Create data rows
    table_data = [headers]
    
    data_style = ParagraphStyle(
        'DataStyle',
        fontSize=9,
        alignment=0,  # LEFT
    )
    
    # Recurring row
    recurring_row = [
        'Recurring',
        f"{float(thub_data.get('total_funds_released', 0)):,.2f}",
        f"{float(thub_data.get('total_expenditure', 0)):,.2f}",
        f"{float(thub_data.get('balance', 0)):,.2f}",
        ''
    ]
    table_data.append(recurring_row)
    
    # Total row
    total_row = [
        'Total',
        f"{float(thub_data.get('total_funds_released', 0)):,.2f}",
        f"{float(thub_data.get('total_expenditure', 0)):,.2f}",
        f"{float(thub_data.get('balance', 0)):,.2f}",
        ''
    ]
    table_data.append(total_row)
    
    # Create table with proper column widths
    table = Table(table_data, colWidths=[1.2*inch, 1.5*inch, 1.5*inch, 1.8*inch, 1.0*inch])
    table.setStyle(TableStyle([
        # Header row styling - white background with bold dark text
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFFFFF')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#000000')),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('LEFTPADDING', (0, 0), (-1, 0), 5),
        ('RIGHTPADDING', (0, 0), (-1, 0), 5),
        
        # Data rows styling
        ('FONTNAME', (0, 1), (-1, 1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
        ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('LEFTPADDING', (0, 1), (-1, -1), 5),
        ('RIGHTPADDING', (0, 1), (-1, -1), 5),
        
        # Total row styling (last row)
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#E3F2FD')),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        
        # Grid and borders
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CCCCCC')),
        ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#000000')),
    ]))
    
    elements.append(table)
    
    doc.build(elements)
    output.seek(0)
    return output

def generate_thub_totals_word(thub_data, to_date):
    """Generate T-Hub Totals in Word format"""
    doc = Document()
    
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("T-Hub-Wise Expenditure Summary")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    # Add table
    table = doc.add_table(rows=3, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    headers = ['Sanctioned Head (I)', 'Total Funds Released (II)', 'Total Expenditure (III)', 
               f'Balance as on ({to_date}) (IV = II - III)', 'Remarks (if any)']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data rows
    recurring_cells = table.rows[1].cells
    recurring_cells[0].text = 'Recurring'
    recurring_cells[1].text = f"{float(thub_data.get('total_funds_released', 0)):,.2f}"
    recurring_cells[2].text = f"{float(thub_data.get('total_expenditure', 0)):,.2f}"
    recurring_cells[3].text = f"{float(thub_data.get('balance', 0)):,.2f}"
    recurring_cells[4].text = ''
    
    total_cells = table.rows[2].cells
    total_cells[0].text = 'Total'
    total_cells[1].text = f"{float(thub_data.get('total_funds_released', 0)):,.2f}"
    total_cells[2].text = f"{float(thub_data.get('total_expenditure', 0)):,.2f}"
    total_cells[3].text = f"{float(thub_data.get('balance', 0)):,.2f}"
    total_cells[4].text = ''
    
    # Make total row bold
    for cell in total_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

@app.route('/download-thub-tgs-comparison', methods=['POST'])
def download_thub_tgs_comparison():
    """Download T-Hub & TGs Comparison in selected format (Excel, PDF, or Word)"""
    try:
        data = request.get_json()
        format_type = data.get('format', 'excel')
        comparison_data = data.get('data', {})
        
        if not comparison_data:
            return jsonify({'error': 'No comparison data provided'}), 400
        
        if format_type == 'excel':
            output = generate_thub_tgs_comparison_excel(comparison_data)
            filename = 'T-Hub_TGs_Comparison.xlsx'
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif format_type == 'pdf':
            output = generate_thub_tgs_comparison_pdf(comparison_data)
            filename = 'T-Hub_TGs_Comparison.pdf'
            mimetype = 'application/pdf'
        elif format_type == 'word':
            output = generate_thub_tgs_comparison_word(comparison_data)
            filename = 'T-Hub_TGs_Comparison.docx'
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            return jsonify({'error': 'Invalid format'}), 400
        
        return send_file(
            output,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        print(f"Error downloading T-Hub & TGs comparison: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

def generate_thub_tgs_comparison_excel(comparison_data):
    """Generate T-Hub & TGs Comparison in Excel format"""
    wb = Workbook()
    ws = wb.active
    ws.title = "T-Hub & TGs Comparison"
    
    # Add title
    ws['A1'] = "Total Expenditure (T-Hub & TGs)"
    ws['A1'].font = Font(bold=True, size=14)
    
    # Add headers
    headers = ['T-Hub & TG (I)', 'Total Funds Released (II)', 'Total Expenditure (III)', 
               f'Balance as on ({comparison_data.get("toDate", "DD/MM/YYYY")}) (IV = II - III)', 'Remarks (if any)']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col)
        cell.value = header
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="343A40", end_color="343A40", fill_type="solid")
    
    # Add data rows
    thub_row = [
        'T-Hub',
        comparison_data.get('thubFundsReleased', 0),
        comparison_data.get('thubExpenditure', 0),
        comparison_data.get('thubBalance', 0),
        ''
    ]
    
    tgs_row = [
        'TGs',
        comparison_data.get('tgsFundsReleased', 0),
        comparison_data.get('tgsExpenditure', 0),
        comparison_data.get('tgsBalance', 0),
        ''
    ]
    
    total_row = [
        'Total',
        (comparison_data.get('thubFundsReleased', 0) + comparison_data.get('tgsFundsReleased', 0)),
        (comparison_data.get('thubExpenditure', 0) + comparison_data.get('tgsExpenditure', 0)),
        (comparison_data.get('thubBalance', 0) + comparison_data.get('tgsBalance', 0)),
        ''
    ]
    
    for row_idx, row_data in enumerate([thub_row, tgs_row, total_row], 4):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            if col_idx == 1:
                cell.value = value
            else:
                if isinstance(value, (int, float)) and value != '':
                    cell.value = float(value)
                    cell.number_format = '#,##0.00'
                else:
                    cell.value = value
    
    # Format Total row
    for col in range(1, 6):
        ws.cell(row=6, column=col).font = Font(bold=True)
        ws.cell(row=6, column=col).fill = PatternFill(start_color="E3F2FD", end_color="E3F2FD", fill_type="solid")
    
    # Set column widths
    ws.column_dimensions['A'].width = 20
    ws.column_dimensions['B'].width = 25
    ws.column_dimensions['C'].width = 25
    ws.column_dimensions['D'].width = 30
    ws.column_dimensions['E'].width = 20
    
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def generate_thub_tgs_comparison_pdf(comparison_data):
    """Generate T-Hub & TGs Comparison in PDF format"""
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch, 
                           leftMargin=0.5*inch, rightMargin=0.5*inch)
    elements = []
    
    # Add title
    title_style = ParagraphStyle(
        'Title',
        fontSize=16,
        textColor=colors.HexColor('#0066CC'),
        spaceAfter=20,
        alignment=1,  # CENTER
        fontName='Helvetica-Bold'
    )
    title = Paragraph("Total Expenditure (T-Hub & TGs)", title_style)
    elements.append(title)
    
    # Create header style
    header_style = ParagraphStyle(
        'HeaderStyle',
        fontSize=9,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#000000'),
        alignment=1,  # CENTER
        wordWrap='CJK',
        leading=12
    )
    
    # Create headers
    headers = [
        Paragraph('T-Hub & TG<br/>(I)', header_style),
        Paragraph('Total Funds Released<br/>(II)', header_style),
        Paragraph('Total Expenditure<br/>(III)', header_style),
        Paragraph(f'Balance as on ({comparison_data.get("toDate", "DD/MM/YYYY")})<br/>(IV = II - III)', header_style),
        Paragraph('Remarks<br/>(if any)', header_style)
    ]
    
    # Create data rows
    table_data = [headers]
    
    thub_row = [
        'T-Hub',
        f"{float(comparison_data.get('thubFundsReleased', 0)):,.2f}",
        f"{float(comparison_data.get('thubExpenditure', 0)):,.2f}",
        f"{float(comparison_data.get('thubBalance', 0)):,.2f}",
        ''
    ]
    
    tgs_row = [
        'TGs',
        f"{float(comparison_data.get('tgsFundsReleased', 0)):,.2f}",
        f"{float(comparison_data.get('tgsExpenditure', 0)):,.2f}",
        f"{float(comparison_data.get('tgsBalance', 0)):,.2f}",
        ''
    ]
    
    total_row = [
        'Total',
        f"{float(comparison_data.get('thubFundsReleased', 0) + comparison_data.get('tgsFundsReleased', 0)):,.2f}",
        f"{float(comparison_data.get('thubExpenditure', 0) + comparison_data.get('tgsExpenditure', 0)):,.2f}",
        f"{float(comparison_data.get('thubBalance', 0) + comparison_data.get('tgsBalance', 0)):,.2f}",
        ''
    ]
    
    table_data.append(thub_row)
    table_data.append(tgs_row)
    table_data.append(total_row)
    
    # Create table
    table = Table(table_data, colWidths=[1.2*inch, 1.5*inch, 1.5*inch, 1.8*inch, 1.0*inch])
    table.setStyle(TableStyle([
        # Header row styling
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFFFFF')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#000000')),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('LEFTPADDING', (0, 0), (-1, 0), 5),
        ('RIGHTPADDING', (0, 0), (-1, 0), 5),
        
        # Data rows styling
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
        ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('LEFTPADDING', (0, 1), (-1, -1), 5),
        ('RIGHTPADDING', (0, 1), (-1, -1), 5),
        
        # Total row styling
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#E3F2FD')),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        
        # Grid and borders
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#CCCCCC')),
        ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#000000')),
    ]))
    
    elements.append(table)
    
    doc.build(elements)
    output.seek(0)
    return output

def generate_thub_tgs_comparison_word(comparison_data):
    """Generate T-Hub & TGs Comparison in Word format"""
    doc = Document()
    
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("Total Expenditure (T-Hub & TGs)")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    # Add table
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    headers = ['T-Hub & TG (I)', 'Total Funds Released (II)', 'Total Expenditure (III)', 
               f'Balance as on ({comparison_data.get("toDate", "DD/MM/YYYY")}) (IV = II - III)', 'Remarks (if any)']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add T-Hub row
    thub_cells = table.rows[1].cells
    thub_cells[0].text = 'T-Hub'
    thub_cells[1].text = f"{float(comparison_data.get('thubFundsReleased', 0)):,.2f}"
    thub_cells[2].text = f"{float(comparison_data.get('thubExpenditure', 0)):,.2f}"
    thub_cells[3].text = f"{float(comparison_data.get('thubBalance', 0)):,.2f}"
    thub_cells[4].text = ''
    
    # Add TGs row
    tgs_cells = table.rows[2].cells
    tgs_cells[0].text = 'TGs'
    tgs_cells[1].text = f"{float(comparison_data.get('tgsFundsReleased', 0)):,.2f}"
    tgs_cells[2].text = f"{float(comparison_data.get('tgsExpenditure', 0)):,.2f}"
    tgs_cells[3].text = f"{float(comparison_data.get('tgsBalance', 0)):,.2f}"
    tgs_cells[4].text = ''
    
    # Add Total row
    total_cells = table.rows[3].cells
    total_cells[0].text = 'Total'
    total_cells[1].text = f"{float(comparison_data.get('thubFundsReleased', 0) + comparison_data.get('tgsFundsReleased', 0)):,.2f}"
    total_cells[2].text = f"{float(comparison_data.get('thubExpenditure', 0) + comparison_data.get('tgsExpenditure', 0)):,.2f}"
    total_cells[3].text = f"{float(comparison_data.get('thubBalance', 0) + comparison_data.get('tgsBalance', 0)):,.2f}"
    total_cells[4].text = ''
    
    # Make all total row cells bold
    for cell in total_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

@app.route('/download-all-documents', methods=['POST'])
def download_all_documents():
    """Download all document tables in selected format (Excel, PDF, or Word)"""
    try:
        data = request.get_json()
        format_type = data.get('format', 'excel')
        thub_tgs_data = data.get('thubTgs', {})
        thub_totals_data = data.get('thubTotals', {})
        tg_details_data = data.get('tgDetails', [])
        
        # Debug logging to see what data is received
        print(f"DEBUG: Received thub_tgs_data: {thub_tgs_data}")
        print(f"DEBUG: Received thub_totals_data: {thub_totals_data}")
        print(f"DEBUG: Format type: {format_type}")
        
        if format_type == 'excel':
            output = generate_all_documents_excel(thub_tgs_data, thub_totals_data, tg_details_data)
            filename = 'All_Document_Tables.xlsx'
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif format_type == 'pdf':
            output = generate_all_documents_pdf(thub_tgs_data, thub_totals_data, tg_details_data)
            filename = 'All_Document_Tables.pdf'
            mimetype = 'application/pdf'
        elif format_type == 'word':
            output = generate_all_documents_word(thub_tgs_data, thub_totals_data, tg_details_data)
            filename = 'All_Document_Tables.docx'
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            return jsonify({'error': 'Invalid format'}), 400
        
        return send_file(
            output,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        print(f"Error downloading all documents: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

def generate_all_documents_excel(thub_tgs_data, thub_totals_data, tg_details_data):
    """Generate all document tables in Excel format"""
    wb = Workbook()
    wb.remove(wb.active)
    
    # Sheet 1: T-Hub & TGs Comparison
    ws1 = wb.create_sheet("T-Hub & TGs Comparison", 0)
    ws1['A1'] = "Total Expenditure (T-Hub & TGs)"
    ws1['A1'].font = Font(bold=True, size=14)
    
    headers = ['T-Hub & TG (I)', 'Total Funds Released (II)', 'Total Expenditure (III)', 
               f'Balance as on ({thub_tgs_data.get("toDate", "DD/MM/YYYY")}) (IV = II - III)', 'Remarks (if any)']
    for col, header in enumerate(headers, 1):
        cell = ws1.cell(row=3, column=col)
        cell.value = header
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="343A40", end_color="343A40", fill_type="solid")
    
    row_num = 4
    # Always show T-Hub row if there's data (check for non-None, not zero)
    thub_funds = thub_tgs_data.get('thubFundsReleased')
    thub_exp = thub_tgs_data.get('thubExpenditure')
    thub_bal = thub_tgs_data.get('thubBalance')
    if thub_funds is not None or thub_exp is not None or thub_bal is not None:
        ws1.cell(row=row_num, column=1).value = 'T-Hub'
        ws1.cell(row=row_num, column=2).value = thub_funds if thub_funds is not None else 0
        ws1.cell(row=row_num, column=3).value = thub_exp if thub_exp is not None else 0
        ws1.cell(row=row_num, column=4).value = thub_bal if thub_bal is not None else 0
        row_num += 1
    
    # Always show TGs row if there's data (check for non-None, not zero)
    tgs_funds = thub_tgs_data.get('tgsFundsReleased')
    tgs_exp = thub_tgs_data.get('tgsExpenditure')
    tgs_bal = thub_tgs_data.get('tgsBalance')
    if tgs_funds is not None or tgs_exp is not None or tgs_bal is not None:
        ws1.cell(row=row_num, column=1).value = 'TGs'
        ws1.cell(row=row_num, column=2).value = tgs_funds if tgs_funds is not None else 0
        ws1.cell(row=row_num, column=3).value = tgs_exp if tgs_exp is not None else 0
        ws1.cell(row=row_num, column=4).value = tgs_bal if tgs_bal is not None else 0
        row_num += 1
    
    ws1.cell(row=row_num, column=1).value = 'Total'
    ws1.cell(row=row_num, column=2).value = thub_tgs_data.get('totalFundsReleased', 0)
    ws1.cell(row=row_num, column=3).value = thub_tgs_data.get('totalExpenditure', 0)
    ws1.cell(row=row_num, column=4).value = thub_tgs_data.get('totalBalance', 0)
    
    # Sheet 2: T-Hub-Wise Expenditure Summary
    ws2 = wb.create_sheet("T-Hub-Wise Exp Summary", 1)
    ws2['A1'] = "T-Hub-Wise Expenditure Summary"
    ws2['A1'].font = Font(bold=True, size=14)
    
    headers2 = ['Sanctioned Head (I)', 'Total Funds Released (II)', 'Total Expenditure (III)', 
                f'Balance as on ({thub_totals_data.get("toDate", "DD/MM/YYYY")}) (VI = II - III)', 'Remarks (if any)']
    for col, header in enumerate(headers2, 1):
        cell = ws2.cell(row=3, column=col)
        cell.value = header
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="343A40", end_color="343A40", fill_type="solid")
    
    ws2.cell(row=4, column=1).value = 'Recurring'
    ws2.cell(row=4, column=2).value = thub_totals_data.get('total_funds_released', 0)
    ws2.cell(row=4, column=3).value = thub_totals_data.get('total_expenditure', 0)
    ws2.cell(row=4, column=4).value = thub_totals_data.get('balance', 0)
    
    ws2.cell(row=5, column=1).value = 'Total'
    ws2.cell(row=5, column=2).value = thub_totals_data.get('total_funds_released', 0)
    ws2.cell(row=5, column=3).value = thub_totals_data.get('total_expenditure', 0)
    ws2.cell(row=5, column=4).value = thub_totals_data.get('balance', 0)
    
    # Sheet 3+: TG Details
    for idx, tg_detail in enumerate(tg_details_data):
        ws = wb.create_sheet(f"TG-{idx+1}", idx+2)
        ws['A1'] = tg_detail.get('title', f'TG {idx+1}')
        ws['A1'].font = Font(bold=True, size=14)
        
        if tg_detail.get('rows'):
            headers = tg_detail.get('columns', [])
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=3, column=col)
                cell.value = header
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="343A40", end_color="343A40", fill_type="solid")
            
            for row_idx, row_data in enumerate(tg_detail.get('rows', []), 4):
                for col_idx, value in enumerate(row_data, 1):
                    ws.cell(row=row_idx, column=col_idx).value = value
    
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def generate_all_documents_pdf(thub_tgs_data, thub_totals_data, tg_details_data):
    """Generate all document tables in PDF format with proper A4 alignment"""
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    
    styles = getSampleStyleSheet()
    title_style = styles['Heading2']
    
    # Header cell style for proper text wrapping
    header_cell_style = ParagraphStyle(
        'HeaderCell',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.HexColor('#000000'),
        alignment=TA_CENTER,
        spaceAfter=0,
        leading=10
    )
    
    # Custom style for centered text
    center_style = ParagraphStyle(
        'centered',
        parent=styles['Normal'],
        alignment=TA_CENTER,
        fontSize=9,
        spaceAfter=6
    )
    
    elements = []
    
    # A4 page dimensions: 210mm x 297mm
    # Usable width: ~180mm after margins
    a4_width = 7.08*inch  # Adjusted for A4
    
    # T-Hub & TGs Comparison
    elements.append(Paragraph("Total Expenditure (T-Hub & TGs)", title_style))
    elements.append(Spacer(1, 0.2*inch))
    
    thub_tgs_table_data = [
        [
            Paragraph('T-Hub & TG', header_cell_style),
            Paragraph('Total Funds<br/>Released', header_cell_style),
            Paragraph('Total<br/>Expenditure', header_cell_style),
            Paragraph(f'Balance<br/>({thub_tgs_data.get("toDate", "DD/MM/YYYY")})', header_cell_style),
            Paragraph('Remarks<br/>(if any)', header_cell_style)
        ]
    ]
    if thub_tgs_data.get('thubFundsReleased') is not None or thub_tgs_data.get('thubExpenditure') is not None or thub_tgs_data.get('thubBalance') is not None:
        thub_tgs_table_data.append([
            'T-Hub',
            f"{thub_tgs_data.get('thubFundsReleased', 0):,.2f}",
            f"{thub_tgs_data.get('thubExpenditure', 0):,.2f}",
            f"{thub_tgs_data.get('thubBalance', 0):,.2f}",
            ''
        ])
    if thub_tgs_data.get('tgsFundsReleased') is not None or thub_tgs_data.get('tgsExpenditure') is not None or thub_tgs_data.get('tgsBalance') is not None:
        thub_tgs_table_data.append([
            'TGs',
            f"{thub_tgs_data.get('tgsFundsReleased', 0):,.2f}",
            f"{thub_tgs_data.get('tgsExpenditure', 0):,.2f}",
            f"{thub_tgs_data.get('tgsBalance', 0):,.2f}",
            ''
        ])
    thub_tgs_table_data.append([
        'Total',
        f"{thub_tgs_data.get('totalFundsReleased', 0):,.2f}",
        f"{thub_tgs_data.get('totalExpenditure', 0):,.2f}",
        f"{thub_tgs_data.get('totalBalance', 0):,.2f}",
        ''
    ])
    
    # Create dynamic row heights matching actual table rows (header + data rows)
    table_row_heights = [0.5*inch] + [None] * (len(thub_tgs_table_data) - 1)
    table = Table(thub_tgs_table_data, colWidths=[1.1*inch, 1.2*inch, 1.2*inch, 1.3*inch, 1*inch], rowHeights=table_row_heights)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFFFFF')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#000000')),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 8),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('LEFTPADDING', (0, 0), (-1, 0), 5),
        ('RIGHTPADDING', (0, 0), (-1, 0), 5),
        ('ALIGN', (0, 1), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 7),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 7),
        ('LEFTPADDING', (0, 1), (-1, -1), 5),
        ('RIGHTPADDING', (0, 1), (-1, -1), 5),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')]),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 0.3*inch))
    
    # T-Hub-Wise Expenditure Summary
    elements.append(Paragraph("T-Hub-Wise Expenditure Summary", title_style))
    elements.append(Spacer(1, 0.2*inch))
    
    thub_totals_table_data = [
        [
            Paragraph('Sanctioned Head', header_cell_style),
            Paragraph('Total Funds<br/>Released', header_cell_style),
            Paragraph('Total<br/>Expenditure', header_cell_style),
            Paragraph(f'Balance<br/>({thub_totals_data.get("toDate", "DD/MM/YYYY")})', header_cell_style),
            Paragraph('Remarks<br/>(if any)', header_cell_style)
        ],
        ['Recurring', f"{thub_totals_data.get('total_funds_released', 0):,.2f}",
         f"{thub_totals_data.get('total_expenditure', 0):,.2f}",
         f"{thub_totals_data.get('balance', 0):,.2f}", ''],
        ['Total', f"{thub_totals_data.get('total_funds_released', 0):,.2f}",
         f"{thub_totals_data.get('total_expenditure', 0):,.2f}",
         f"{thub_totals_data.get('balance', 0):,.2f}", '']
    ]
    
    # Create dynamic row heights matching actual table rows (header + data rows)
    table2_row_heights = [0.5*inch] + [None] * (len(thub_totals_table_data) - 1)
    table2 = Table(thub_totals_table_data, colWidths=[1.1*inch, 1.2*inch, 1.2*inch, 1.3*inch, 1*inch], rowHeights=table2_row_heights)
    table2.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFFFFF')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#000000')),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 8),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('LEFTPADDING', (0, 0), (-1, 0), 5),
        ('RIGHTPADDING', (0, 0), (-1, 0), 5),
        ('ALIGN', (0, 1), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 7),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 7),
        ('LEFTPADDING', (0, 1), (-1, -1), 5),
        ('RIGHTPADDING', (0, 1), (-1, -1), 5),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')]),
    ]))
    elements.append(table2)
    elements.append(Spacer(1, 0.15*inch))
    
    # TG Details
    for tg_detail in tg_details_data:
        elements.append(Spacer(1, 0.15*inch))  # Reduced space between sections
        elements.append(Paragraph(tg_detail.get('title', 'TG Details'), title_style))
        elements.append(Spacer(1, 0.1*inch))
        
        if tg_detail.get('rows'):
            # Create header row with Paragraph objects
            columns = tg_detail.get('columns', [])
            tg_table_data = [
                [Paragraph(str(col), header_cell_style) for col in columns]
            ]
            tg_table_data.extend(tg_detail.get('rows', []))
            
            # Calculate dynamic column widths based on number of columns
            num_cols = len(columns)
            col_width = (a4_width - 0.5*inch) / num_cols
            col_widths = [col_width] * num_cols
            
            # Create dynamic row heights matching actual table rows (header + data rows)
            tg_row_heights = [0.45*inch] + [None] * (len(tg_table_data) - 1)
            tg_table = Table(tg_table_data, colWidths=col_widths, rowHeights=tg_row_heights)
            tg_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFFFFF')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#000000')),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 8),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('TOPPADDING', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('LEFTPADDING', (0, 0), (-1, 0), 5),
                ('RIGHTPADDING', (0, 0), (-1, 0), 5),
                ('WORDWRAP', (0, 0), (-1, 0), True),
                ('ALIGN', (0, 1), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                ('LEFTPADDING', (0, 1), (-1, -1), 4),
                ('RIGHTPADDING', (0, 1), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')]),
            ]))
            elements.append(tg_table)
    
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, rightMargin=12*mm, leftMargin=12*mm, topMargin=15*mm, bottomMargin=15*mm)
    doc.build(elements)
    output.seek(0)
    return output

def generate_all_documents_word(thub_tgs_data, thub_totals_data, tg_details_data):
    """Generate all document tables in Word format"""
    doc = Document()
    
    # T-Hub & TGs Comparison
    doc.add_heading('Total Expenditure (T-Hub & TGs)', level=2)
    
    # Count how many data rows we need
    data_rows = 1  # For Total row
    has_thub = thub_tgs_data.get('thubFundsReleased') is not None or thub_tgs_data.get('thubExpenditure') is not None or thub_tgs_data.get('thubBalance') is not None
    has_tgs = thub_tgs_data.get('tgsFundsReleased') is not None or thub_tgs_data.get('tgsExpenditure') is not None or thub_tgs_data.get('tgsBalance') is not None
    if has_thub:
        data_rows += 1
    if has_tgs:
        data_rows += 1
    
    table1 = doc.add_table(rows=data_rows + 1, cols=5)  # +1 for header
    table1.style = 'Table Grid'
    
    headers = ['T-Hub & TG (I)', 'Total Funds Released (II)', 'Total Expenditure (III)',
               f'Balance as on ({thub_tgs_data.get("toDate", "DD/MM/YYYY")}) (IV = II - III)', 'Remarks (if any)']
    for col, header in enumerate(headers):
        table1.rows[0].cells[col].text = header
    
    row_idx = 1
    if has_thub:
        table1.rows[row_idx].cells[0].text = 'T-Hub'
        table1.rows[row_idx].cells[1].text = str(thub_tgs_data.get('thubFundsReleased', 0))
        table1.rows[row_idx].cells[2].text = str(thub_tgs_data.get('thubExpenditure', 0))
        table1.rows[row_idx].cells[3].text = str(thub_tgs_data.get('thubBalance', 0))
        row_idx += 1
    
    if has_tgs:
        table1.rows[row_idx].cells[0].text = 'TGs'
        table1.rows[row_idx].cells[1].text = str(thub_tgs_data.get('tgsFundsReleased', 0))
        table1.rows[row_idx].cells[2].text = str(thub_tgs_data.get('tgsExpenditure', 0))
        table1.rows[row_idx].cells[3].text = str(thub_tgs_data.get('tgsBalance', 0))
        row_idx += 1
    
    table1.rows[row_idx].cells[0].text = 'Total'
    table1.rows[row_idx].cells[1].text = str(thub_tgs_data.get('totalFundsReleased', 0))
    table1.rows[row_idx].cells[2].text = str(thub_tgs_data.get('totalExpenditure', 0))
    table1.rows[row_idx].cells[3].text = str(thub_tgs_data.get('totalBalance', 0))
    
    doc.add_paragraph()
    
    # T-Hub-Wise Expenditure Summary
    doc.add_heading('T-Hub-Wise Expenditure Summary', level=2)
    
    table2 = doc.add_table(rows=3, cols=5)
    table2.style = 'Table Grid'
    
    headers2 = ['Sanctioned Head (I)', 'Total Funds Released (II)', 'Total Expenditure (III)',
                f'Balance as on ({thub_totals_data.get("toDate", "DD/MM/YYYY")}) (VI = II - III)', 'Remarks (if any)']
    for col, header in enumerate(headers2):
        table2.rows[0].cells[col].text = header
    
    table2.rows[1].cells[0].text = 'Recurring'
    table2.rows[1].cells[1].text = str(thub_totals_data.get('total_funds_released', 0))
    table2.rows[1].cells[2].text = str(thub_totals_data.get('total_expenditure', 0))
    table2.rows[1].cells[3].text = str(thub_totals_data.get('balance', 0))
    
    table2.rows[2].cells[0].text = 'Total'
    table2.rows[2].cells[1].text = str(thub_totals_data.get('total_funds_released', 0))
    table2.rows[2].cells[2].text = str(thub_totals_data.get('total_expenditure', 0))
    table2.rows[2].cells[3].text = str(thub_totals_data.get('balance', 0))
    
    doc.add_paragraph()
    
    # TG Details
    for tg_detail in tg_details_data:
        doc.add_heading(tg_detail.get('title', 'TG Details'), level=2)
        
        if tg_detail.get('rows'):
            num_cols = len(tg_detail.get('columns', []))
            table_tg = doc.add_table(rows=len(tg_detail.get('rows', [])) + 1, cols=num_cols)
            table_tg.style = 'Table Grid'
            
            for col, header in enumerate(tg_detail.get('columns', [])):
                table_tg.rows[0].cells[col].text = str(header)
            
            for row_idx, row_data in enumerate(tg_detail.get('rows', []), 1):
                for col_idx, value in enumerate(row_data):
                    table_tg.rows[row_idx].cells[col_idx].text = str(value)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)